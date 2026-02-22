"""
Test document export with proper mock data matching prompt structure
"""
import requests
import json
from pathlib import Path

BASE_URL = "http://localhost:8000/api/v1"

# Correct BA structure (matching brd_prompts_optimized.py examples)
CORRECT_BA_MOCK = {
    "ekranlar": [
        {
            "ekran_adi": "Login Ekranı",
            "aciklama": "Kullanıcıların email ve şifre ile sisteme giriş yapmasını sağlar. Şifre sıfırlama işlevini destekler.",
            "gerekli_dokumanlar": {
                "teknik_akis": "Login akış diyagramı gereklidir",
                "tasarim_dosyasi": "Figma - Login Screen Design v2.0"
            },
            "is_akisi_diyagrami": [
                "1. Kullanıcı login sayfasına gider",
                "2. Email adresini girer",
                "3. Şifresini girer",
                "4. Giriş Yap butonuna tıklar",
                "5. Sistem bilgileri doğrular",
                "6. Başarılı ise ana sayfaya yönlendirilir",
                "7. Hatalı ise uyarı mesajı gösterilir"
            ],
            "fonksiyonel_gereksinimler": [
                {
                    "id": "FR-01",
                    "tanim": "Kullanıcı email adresi girebilmelidir"
                },
                {
                    "id": "FR-02",
                    "tanim": "Kullanıcı şifre girebilmelidir"
                },
                {
                    "id": "FR-03",
                    "tanim": "Giriş Yap butonu tıklanabilir olmalıdır"
                },
                {
                    "id": "FR-04",
                    "tanim": "Şifremi Unuttum linki görünür ve tıklanabilir olmalıdır"
                },
                {
                    "id": "FR-05",
                    "tanim": "Hatalı giriş denemesinde kullanıcıya bilgilendirici hata mesajı gösterilmelidir"
                }
            ],
            "is_kurallari": [
                {
                    "kural": "3 başarısız giriş denemesinden sonra hesap 15 dakika bloklanır",
                    "detay": "Brute force saldırılarını önlemek için"
                },
                {
                    "kural": "Şifre minimum 8 karakter olmalıdır",
                    "detay": "Güvenlik standardı gereği"
                },
                {
                    "kural": "Email formatı geçerli olmalıdır",
                    "detay": "@  ve domain içermelidir"
                }
            ],
            "kabul_kriterleri": [
                {
                    "id": "BR-01",
                    "kriter": "Geçerli email ve şifre ile giriş başarılı olmalıdır"
                },
                {
                    "id": "BR-02",
                    "kriter": "Hatalı email veya şifre ile giriş engellenmelidir"
                },
                {
                    "id": "BR-03",
                    "kriter": "3 hatalı denemeden sonra 'Hesap bloklandı' mesajı gösterilmelidir"
                },
                {
                    "id": "BR-04",
                    "kriter": "Şifremi Unuttum linki şifre sıfırlama sayfasına yönlendirmelidir"
                }
            ],
            "validasyonlar": [
                {
                    "alan": "email",
                    "kisit": "Email formatı (@domain.com)",
                    "hata_mesaji": "Geçerli bir email adresi giriniz"
                },
                {
                    "alan": "email",
                    "kisit": "Zorunlu alan",
                    "hata_mesaji": "Email adresi boş bırakılamaz"
                },
                {
                    "alan": "password",
                    "kisit": "Minimum 8 karakter",
                    "hata_mesaji": "Şifre en az 8 karakter olmalıdır"
                },
                {
                    "alan": "password",
                    "kisit": "Zorunlu alan",
                    "hata_mesaji": "Şifre boş bırakılamaz"
                }
            ]
        },
        {
            "ekran_adi": "Dashboard",
            "aciklama": "Ana kontrol paneli. Kullanıcı istatistiklerini ve son aktiviteleri gösterir.",
            "gerekli_dokumanlar": {
                "teknik_akis": "Dashboard data flow diagram",
                "tasarim_dosyasi": "Figma - Dashboard v1.5"
            },
            "is_akisi_diyagrami": [
                "1. Kullanıcı başarılı giriş yapar",
                "2. Dashboard sayfası yüklenir",
                "3. Kullanıcı istatistikleri gösterilir",
                "4. Son aktiviteler listelenir"
            ],
            "fonksiyonel_gereksinimler": [
                {
                    "id": "FR-06",
                    "tanim": "Kullanıcı adı ve profil resmi görüntülenmelidir"
                },
                {
                    "id": "FR-07",
                    "tanim": "Son 5 aktivite listelenmelidir"
                },
                {
                    "id": "FR-08",
                    "tanim": "İstatistik kartları (toplam görev, tamamlanan, bekleyen) gösterilmelidir"
                }
            ],
            "is_kurallari": [
                {
                    "kural": "Sadece giriş yapmış kullanıcılar erişebilir",
                    "detay": "Authentication token kontrolü yapılır"
                }
            ],
            "kabul_kriterleri": [
                {
                    "id": "BR-05",
                    "kriter": "Dashboard başarılı giriş sonrası otomatik yüklenmelidir"
                },
                {
                    "id": "BR-06",
                    "kriter": "Tüm istatistikler gerçek zamanlı güncellenmelidir"
                }
            ],
            "validasyonlar": []
        }
    ]
}

# Correct TA structure (matching Streamlit build_ta_docx)
CORRECT_TA_MOCK = {
    "teknik_analiz": {
        "genel_tanim": {
            "modul_adi": "Kullanıcı Authentication Modülü",
            "teknoloji_stack": ["React 18", "FastAPI", "PostgreSQL", "Redis", "JWT"],
            "mimari_yaklasim": "Microservices with REST API, JWT authentication, Redis session storage"
        },
        "endpoint_detaylari": {
            "/api/v1/auth/login": {
                "method": "POST",
                "aciklama": "Kullanıcı giriş endpoint'i. Email ve şifre alır, JWT token döner.",
                "request_body": {
                    "email": "string",
                    "password": "string",
                    "remember_me": "boolean (optional)"
                },
                "response_success": {
                    "access_token": "string (JWT)",
                    "refresh_token": "string",
                    "user": {
                        "id": "uuid",
                        "email": "string",
                        "name": "string"
                    }
                },
                "response_errors": [
                    {
                        "http_code": "401",
                        "error_code": "INVALID_CREDENTIALS",
                        "mesaj": "Email veya şifre hatalı"
                    },
                    {
                        "http_code": "423",
                        "error_code": "ACCOUNT_LOCKED",
                        "mesaj": "Hesap geçici olarak kilitlendi. 15 dakika sonra tekrar deneyin."
                    }
                ]
            },
            "/api/v1/auth/logout": {
                "method": "POST",
                "aciklama": "Kullanıcı çıkış yapar, session invalidate edilir",
                "request_body": {},
                "response_success": {
                    "message": "Logout successful"
                },
                "response_errors": []
            }
        },
        "dto_veri_yapilari": [
            {
                "dto_adi": "LoginRequest",
                "aciklama": "Login endpoint'i için request DTO",
                "fields": [
                    {
                        "field": "email",
                        "tip": "string",
                        "validasyon": "required, email format"
                    },
                    {
                        "field": "password",
                        "tip": "string",
                        "validasyon": "required, min 8 chars"
                    },
                    {
                        "field": "remember_me",
                        "tip": "boolean",
                        "validasyon": "optional, default false"
                    }
                ]
            },
            {
                "dto_adi": "AuthResponse",
                "aciklama": "Authentication başarılı response",
                "fields": [
                    {
                        "field": "access_token",
                        "tip": "string (JWT)",
                        "validasyon": "required"
                    },
                    {
                        "field": "refresh_token",
                        "tip": "string",
                        "validasyon": "required"
                    },
                    {
                        "field": "user",
                        "tip": "UserDTO",
                        "validasyon": "required"
                    }
                ]
            }
        ],
        "validasyon_kurallari": [
            {
                "id": "VR-01",
                "field": "email",
                "kural": "Email RFC 5322 standardına uygun olmalı"
            },
            {
                "id": "VR-02",
                "field": "password",
                "kural": "Minimum 8 karakter, en az 1 büyük harf, 1 küçük harf, 1 rakam"
            },
            {
                "id": "VR-03",
                "field": "jwt_token",
                "kural": "Token expiration kontrolü yapılmalı"
            }
        ],
        "mock_curl_ornekleri": [
            {
                "endpoint_adi": "POST /api/v1/auth/login",
                "curl": 'curl -X POST http://localhost:8000/api/v1/auth/login \\\n  -H "Content-Type: application/json" \\\n  -d \'{"email": "test@example.com", "password": "Test1234"}\''
            },
            {
                "endpoint_adi": "POST /api/v1/auth/logout",
                "curl": 'curl -X POST http://localhost:8000/api/v1/auth/logout \\\n  -H "Authorization: Bearer <access_token>"'
            }
        ]
    }
}

# Correct TC structure (matching test case format)
CORRECT_TC_MOCK = {
    "test_cases": [
        {
            "test_case_id": "TC-001",
            "br_id": "BR-01",
            "priority": "High",
            "test_area": "Authentication",
            "testcase": "Geçerli email ve şifre ile başarılı giriş",
            "test_steps": "1. Login sayfasına git\n2. Email alanına 'test@example.com' gir\n3. Şifre alanına 'Test1234' gir\n4. Giriş Yap butonuna tıkla\n5. Dashboard'ın yüklendiğini doğrula",
            "test_data": '{"email": "test@example.com", "password": "Test1234"}',
            "expected_result": "Kullanıcı başarıyla giriş yapar ve Dashboard sayfasına yönlendirilir. Access token ve user bilgileri döner."
        },
        {
            "test_case_id": "TC-002",
            "br_id": "BR-02",
            "priority": "High",
            "test_area": "Authentication",
            "testcase": "Hatalı şifre ile giriş denemesi",
            "test_steps": "1. Login sayfasına git\n2. Email alanına 'test@example.com' gir\n3. Şifre alanına 'WrongPassword' gir\n4. Giriş Yap butonuna tıkla\n5. Hata mesajını doğrula",
            "test_data": '{"email": "test@example.com", "password": "WrongPassword"}',
            "expected_result": "Sistem 401 hatası döner. 'Email veya şifre hatalı' mesajı gösterilir. Kullanıcı login sayfasında kalır."
        },
        {
            "test_case_id": "TC-003",
            "br_id": "BR-03",
            "priority": "High",
            "test_area": "Security",
            "testcase": "3 başarısız denemeden sonra hesap kilitleme",
            "test_steps": "1. Login sayfasına git\n2. 3 kez hatalı şifre ile giriş dene\n3. 4. denemede hesap kilitleme mesajını doğrula\n4. 15 dakika bekle\n5. Geçerli şifre ile giriş yapmayı dene",
            "test_data": '{"email": "test@example.com", "wrong_attempts": 3}',
            "expected_result": "3. denemeden sonra 423 hatası döner. 'Hesap geçici olarak kilitlendi' mesajı gösterilir. 15 dakika sonra normal giriş yapılabilir."
        },
        {
            "test_case_id": "TC-004",
            "br_id": "BR-01",
            "priority": "Medium",
            "test_area": "Authentication",
            "testcase": "Email formatı validasyonu",
            "test_steps": "1. Login sayfasına git\n2. Email alanına geçersiz format gir (örn: 'notanemail')\n3. Şifre gir\n4. Giriş Yap'a tıkla\n5. Validasyon hatasını doğrula",
            "test_data": '{"email": "notanemail", "password": "Test1234"}',
            "expected_result": "Frontend validasyonu devreye girer. 'Geçerli bir email adresi giriniz' mesajı gösterilir. Backend'e istek gönderilmez."
        },
        {
            "test_case_id": "TC-005",
            "br_id": "BR-04",
            "priority": "Medium",
            "test_area": "Password Reset",
            "testcase": "Şifremi Unuttum linki çalışması",
            "test_steps": "1. Login sayfasına git\n2. 'Şifremi Unuttum' linkine tıkla\n3. Şifre sıfırlama sayfasına yönlendirildiğini doğrula",
            "test_data": "{}",
            "expected_result": "Kullanıcı /forgot-password sayfasına yönlendirilir. Email girme formu gösterilir."
        }
    ]
}


def test_with_mock_data():
    print("\n" + "="*80)
    print("TESTING DOCUMENT EXPORT WITH CORRECT MOCK DATA")
    print("="*80)

    # Get existing project
    print("\n[1/5] Getting existing project...")
    projects_response = requests.get(f"{BASE_URL}/projects")
    if projects_response.status_code != 200:
        print(f"✗ Failed to get projects: {projects_response.status_code}")
        return

    projects = projects_response.json()
    if not projects:
        print("✗ No projects found. Please create a project first.")
        return

    project_id = projects[0]["id"]
    project_name = projects[0]["name"]
    print(f"✓ Using project: {project_name} (ID={project_id})")

    # Start pipeline
    print("\n[2/5] Starting pipeline...")
    pipeline_response = requests.post(f"{BASE_URL}/pipeline/start", json={
        "project_id": project_id,
        "project_name": project_name,
        "brd_content": "# Test BRD\n\nLogin ve Dashboard ekranları",
        "stages": ["ba", "ta", "tc"]
    })

    if pipeline_response.status_code != 202:
        print(f"✗ Failed to start pipeline: {pipeline_response.status_code}")
        return

    run_id = pipeline_response.json()["pipeline_run_id"]
    print(f"✓ Pipeline started: run_id={run_id}")

    # Save correct mock data
    print("\n[3/5] Saving correct mock data...")

    # BA
    ba_response = requests.post(
        f"{BASE_URL}/pipeline/{run_id}/checkpoint/ba_gen",
        json=CORRECT_BA_MOCK
    )
    if ba_response.status_code != 200:
        print(f"✗ BA checkpoint failed: {ba_response.status_code}")
        print(ba_response.text)
        return
    print("✓ BA checkpoint saved")

    # TA
    ta_response = requests.post(
        f"{BASE_URL}/pipeline/{run_id}/checkpoint/ta_gen",
        json=CORRECT_TA_MOCK
    )
    if ta_response.status_code != 200:
        print(f"✗ TA checkpoint failed: {ta_response.status_code}")
        return
    print("✓ TA checkpoint saved")

    # TC
    tc_response = requests.post(
        f"{BASE_URL}/pipeline/{run_id}/checkpoint/tc_gen",
        json=CORRECT_TC_MOCK
    )
    if tc_response.status_code != 200:
        print(f"✗ TC checkpoint failed: {tc_response.status_code}")
        return
    print("✓ TC checkpoint saved")

    # Test downloads
    print("\n[4/5] Testing downloads...")
    output_dir = Path("/tmp/ba_qa_exports_correct")
    output_dir.mkdir(exist_ok=True)

    for doc_type in ['ba', 'ta', 'tc']:
        download_response = requests.get(
            f"{BASE_URL}/documents-export/{run_id}/download/{doc_type}"
        )

        if download_response.status_code == 200:
            extension = 'xlsx' if doc_type == 'tc' else 'docx'
            filename = output_dir / f"{doc_type.upper()}_correct.{extension}"

            with open(filename, 'wb') as f:
                f.write(download_response.content)

            print(f"  ✓ {doc_type.upper()} - {filename} ({len(download_response.content)} bytes)")
        else:
            print(f"  ✗ {doc_type.upper()} failed: {download_response.status_code}")

    # Verify DOCX structure
    print("\n[5/5] Verifying document structure...")
    try:
        from docx import Document

        # Check BA
        ba_doc = Document(output_dir / "BA_correct.docx")
        ba_paragraphs = len(ba_doc.paragraphs)
        ba_headings = len([p for p in ba_doc.paragraphs if p.style.name.startswith('Heading')])
        print(f"  ✓ BA: {ba_paragraphs} paragraphs, {ba_headings} headings")

        # Check TA
        ta_doc = Document(output_dir / "TA_correct.docx")
        ta_paragraphs = len(ta_doc.paragraphs)
        ta_headings = len([p for p in ta_doc.paragraphs if p.style.name.startswith('Heading')])
        print(f"  ✓ TA: {ta_paragraphs} paragraphs, {ta_headings} headings")

        # Check TC
        from openpyxl import load_workbook
        tc_wb = load_workbook(output_dir / "TC_correct.xlsx")
        tc_ws = tc_wb.active
        tc_rows = tc_ws.max_row
        print(f"  ✓ TC: {tc_rows} rows (including header)")

    except Exception as e:
        print(f"  ⚠ Verification error: {e}")

    print("\n" + "="*80)
    print("🎉 TEST COMPLETE WITH CORRECT MOCK DATA!")
    print("="*80)
    print(f"\nExported files: {output_dir}")
    print("\nOpen the files to verify the correct format:")
    print(f"  - open {output_dir}/BA_correct.docx")
    print(f"  - open {output_dir}/TA_correct.docx")
    print(f"  - open {output_dir}/TC_correct.xlsx")


if __name__ == "__main__":
    test_with_mock_data()
