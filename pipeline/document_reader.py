"""
Document Reader - Extract text from Word documents and Google Drive
"""
import io
import re
import tempfile
from typing import Optional, Tuple, List, Dict, Any
import requests


def read_docx(file_content: bytes) -> str:
    """
    Extract text from a Word (.docx) document

    Args:
        file_content: Binary content of the .docx file

    Returns:
        Extracted text with structure preserved
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx library is required. Install with: pip install python-docx")

    # Load document from bytes
    doc = Document(io.BytesIO(file_content))

    # Extract text with structure
    text_parts = []

    for paragraph in doc.paragraphs:
        # Preserve heading structure
        style = paragraph.style.name.lower()
        text = paragraph.text.strip()

        if not text:
            text_parts.append("")  # Preserve empty lines
            continue

        # Detect heading levels
        if 'heading 1' in style or 'title' in style:
            text_parts.append(f"# {text}")
        elif 'heading 2' in style:
            text_parts.append(f"## {text}")
        elif 'heading 3' in style:
            text_parts.append(f"### {text}")
        else:
            # Regular paragraph
            text_parts.append(text)

    return "\n".join(text_parts)


def extract_google_drive_file_id(url: str) -> Optional[str]:
    """
    Extract file ID from Google Drive URL

    Supports formats:
    - https://drive.google.com/file/d/FILE_ID/view
    - https://drive.google.com/open?id=FILE_ID
    - https://docs.google.com/document/d/FILE_ID/edit
    - https://docs.google.com/spreadsheets/d/FILE_ID/edit

    Args:
        url: Google Drive or Google Docs URL

    Returns:
        File ID or None if not found
    """
    # Pattern 1: /d/FILE_ID/ (most common)
    match = re.search(r'/d/([a-zA-Z0-9_-]+)', url)
    if match:
        return match.group(1)

    # Pattern 2: ?id=FILE_ID
    match = re.search(r'[?&]id=([a-zA-Z0-9_-]+)', url)
    if match:
        return match.group(1)

    return None


def download_from_google_drive(url: str) -> Tuple[bytes, str]:
    """
    Download document from Google Drive

    Args:
        url: Google Drive share link (must be publicly accessible or "Anyone with link")

    Returns:
        Tuple of (file_content, mime_type)

    Raises:
        ValueError: If URL is invalid or file ID cannot be extracted
        requests.RequestException: If download fails
    """
    # Extract file ID
    file_id = extract_google_drive_file_id(url)
    if not file_id:
        raise ValueError("Could not extract file ID from Google Drive URL. "
                        "Make sure it's a valid Google Drive share link.")

    # Try to download directly (works for publicly shared files)
    download_url = f"https://drive.google.com/uc?export=download&id={file_id}"

    try:
        response = requests.get(download_url, timeout=30)
        response.raise_for_status()

        # Check if we got a download confirmation page (for large files)
        if 'download_warning' in response.text or 'virus scan warning' in response.text.lower():
            # Extract confirm token
            confirm_match = re.search(r'confirm=([0-9A-Za-z_-]+)', response.text)
            if confirm_match:
                confirm = confirm_match.group(1)
                download_url = f"https://drive.google.com/uc?export=download&confirm={confirm}&id={file_id}"
                response = requests.get(download_url, timeout=30)
                response.raise_for_status()

        # Detect content type
        content_type = response.headers.get('content-type', 'application/octet-stream')

        # Check if we actually got file content (not HTML error page)
        if 'text/html' in content_type.lower() and len(response.content) < 10000:
            # Likely an error page or permission issue
            if 'Google Drive' in response.text and 'Access' in response.text:
                raise ValueError("Cannot access file. Make sure the Google Drive link is set to "
                               "'Anyone with the link can view' in sharing settings.")
            else:
                raise ValueError("Failed to download file. Make sure the link is a direct file link, "
                               "not a folder or restricted document.")

        return response.content, content_type

    except requests.RequestException as e:
        raise requests.RequestException(f"Failed to download from Google Drive: {str(e)}")


def read_document_from_drive(url: str) -> str:
    """
    Download and extract text from a Google Drive document

    Args:
        url: Google Drive share link

    Returns:
        Extracted text

    Raises:
        ValueError: If URL is invalid or file format is not supported
        requests.RequestException: If download fails
    """
    # Download file
    file_content, mime_type = download_from_google_drive(url)

    # Handle different file types
    if 'word' in mime_type.lower() or 'document' in mime_type.lower() or file_content.startswith(b'PK'):
        # Word document (.docx)
        return read_docx(file_content)

    elif 'text/plain' in mime_type.lower():
        # Plain text
        return file_content.decode('utf-8')

    elif 'application/pdf' in mime_type.lower():
        # PDF - would need additional library
        raise ValueError("PDF files are not yet supported. Please use Word (.docx) or text files.")

    else:
        # Try to detect by content
        if file_content.startswith(b'PK'):
            # Likely a .docx file (ZIP-based)
            return read_docx(file_content)

        # Check if we got HTML (common error when permissions are wrong)
        if file_content.startswith(b'<!DOCTYPE') or file_content.startswith(b'<html'):
            raise ValueError(
                "Received HTML page instead of document content.\n\n"
                "This usually means:\n"
                "1. Document is not publicly accessible\n"
                "2. Sharing is not set to 'Anyone with the link'\n"
                "3. You may need to export as .docx and upload instead\n\n"
                "Please check sharing settings and try again."
            )

        # Try plain text as fallback
        try:
            decoded = file_content.decode('utf-8')

            # Additional check: if decoded text looks like HTML
            if decoded.strip().startswith('<!DOCTYPE') or decoded.strip().startswith('<html'):
                raise ValueError(
                    "Received HTML page instead of document content.\n"
                    "Please set document sharing to 'Anyone with the link can view'."
                )

            return decoded
        except UnicodeDecodeError:
            raise ValueError(f"Unsupported file format (MIME: {mime_type}). "
                           "Please use Word (.docx) or text files.")


def export_google_doc_as_text(doc_url: str) -> str:
    """
    Export a Google Docs document as plain text

    This uses Google Docs export feature for Google Docs native documents

    Args:
        doc_url: Google Docs URL (https://docs.google.com/document/d/...)

    Returns:
        Exported text content
    """
    file_id = extract_google_drive_file_id(doc_url)
    if not file_id:
        raise ValueError("Could not extract document ID from Google Docs URL")

    # Use Google Docs export API
    export_url = f"https://docs.google.com/document/d/{file_id}/export?format=txt"

    try:
        response = requests.get(export_url, timeout=30)
        response.raise_for_status()

        content_type = response.headers.get('content-type', '').lower()
        text = response.text

        # Check if we got HTML instead of text (permission error or login page)
        if 'text/html' in content_type or text.strip().startswith('<!DOCTYPE') or text.strip().startswith('<html'):
            raise ValueError(
                "Cannot access Google Doc - received HTML instead of document content.\n\n"
                "Please check:\n"
                "1. Document sharing is set to 'Anyone with the link can view'\n"
                "2. Link permissions are correct\n"
                "3. Document is not private/restricted\n\n"
                "Alternative: Export document as .docx and upload it instead."
            )

        # Additional check: if response is very large and looks like HTML
        if len(text) > 100000 and '<html' in text.lower()[:1000]:
            raise ValueError(
                "Received HTML page instead of document content.\n\n"
                "This usually means the document is not publicly accessible.\n"
                "Please set sharing to 'Anyone with the link can view'."
            )

        return text

    except requests.RequestException as e:
        # Fallback to regular download method
        return read_document_from_drive(doc_url)


# ---------------------------------------------------------------------------
# Adım 1: Enhanced DOCX Reader — Bullet Level + Hyperlink + Bold
# ---------------------------------------------------------------------------

def read_docx_structured(file_content: bytes) -> List[Dict[str, Any]]:
    """
    DOCX dosyasını yapısal element listesi olarak döndür.
    Body element sırasını korur; heading'leri (H1-H6), liste öğelerini
    (level 0-3), bold segment'leri ve inline hyperlink'leri çıkarır.

    Returns:
        List of elements:
        - {"type": "heading",   "level": 1, "text": "Proje Açıklaması"}
        - {"type": "heading",   "level": 2, "text": "Splash"}
        - {"type": "list_item", "level": 0, "text": "...", "bold_segments": [], "links": []}
        - {"type": "paragraph", "text": "...", "bold_segments": [], "links": []}
        - {"type": "table",     "headers": [...], "rows": [[...], ...]}
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        raise ImportError("python-docx library is required. Install with: pip install python-docx")

    doc = Document(io.BytesIO(file_content))
    rels = doc.part.rels

    # O(1) lookup: XML element → python-docx object
    para_map  = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    elements = []

    for child in doc.element.body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'sdt':
            continue  # Skip TOC block (structured document tag)

        if tag == 'p':
            para = para_map.get(child)
            if not para or not para.text.strip():
                continue
            element = _parse_paragraph_element(para, rels)
            if element:
                elements.append(element)

        elif tag == 'tbl':
            table = table_map.get(child)
            if table:
                elements.append(_parse_table_element(table))

    return elements


def _parse_paragraph_element(para, rels) -> Optional[Dict[str, Any]]:
    """Paragraph'ı type, text, bold_segments ve links bilgisiyle parse et."""
    from docx.oxml.ns import qn

    style = para.style.name.lower() if para.style else ''
    text  = para.text.strip()

    # 1. Heading tespiti (H1-H6)
    for i in range(1, 7):
        if f'heading {i}' in style or (i == 1 and 'title' in style):
            if not text:
                return None
            return {"type": "heading", "level": i, "text": text}

    # 2. List level tespiti (numPr/ilvl XML attribute'ünden)
    list_level = None
    pPr = para._element.find(qn('w:pPr'))
    if pPr is not None:
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            ilvl_el = numPr.find(qn('w:ilvl'))
            if ilvl_el is not None:
                list_level = int(ilvl_el.get(qn('w:val'), '0'))

    # 3. Bold segments (run-level bold detection)
    bold_segments = [r.text for r in para.runs if r.bold and r.text.strip()]

    # 4. Hyperlinks
    links = _extract_hyperlinks_from_para(para._element, rels)

    if list_level is not None:
        return {
            "type":          "list_item",
            "level":         list_level,
            "text":          text,
            "bold_segments": bold_segments,
            "links":         links,
        }

    return {
        "type":          "paragraph",
        "text":          text,
        "bold_segments": bold_segments,
        "links":         links,
    }


def _extract_hyperlinks_from_para(p_element, rels) -> List[str]:
    """Paragraph içindeki w:hyperlink elementlerinden URL'leri çıkar."""
    from docx.oxml.ns import qn
    links = []
    for hyperlink in p_element.findall('.//' + qn('w:hyperlink')):
        # r:id attribute — relationships namespace
        r_id = hyperlink.get(
            '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id'
        )
        if r_id and r_id in rels:
            url = rels[r_id].target_ref
            if url and url not in links:
                links.append(url)
    return links


def _parse_table_element(table) -> Dict[str, Any]:
    """Table'ı headers + rows formatında parse et."""
    rows_data = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_text = '\n'.join(
                p.text.strip() for p in cell.paragraphs if p.text.strip()
            )
            cells.append(cell_text)
        rows_data.append(_deduplicate_merged_cells(cells))

    if not rows_data:
        return {'type': 'table', 'headers': [], 'rows': []}

    return {
        'type':    'table',
        'headers': rows_data[0],
        'rows':    rows_data[1:] if len(rows_data) > 1 else [],
    }


def _deduplicate_merged_cells(row: list) -> list:
    """python-docx'in merged cell tekrarlarını boş string ile değiştir."""
    if not row:
        return row
    result = [row[0]]
    for i in range(1, len(row)):
        result.append('' if row[i] == row[i - 1] else row[i])
    return result
