"""Sayfa 2: Pipeline Sonuç — Görüntüleme & Export.

Export formatları:
- BA: DOCX
- TA: DOCX
- TC: CSV + Excel (23 kolon)
"""
import json
import io
import time
import streamlit as st
import pandas as pd
from components.sidebar import render_custom_sidebar

st.set_page_config(page_title="Pipeline Sonuç", page_icon="📊", layout="wide")
render_custom_sidebar(active_page="pipeline_sonuc")
st.title("📊 Pipeline Sonuçları")

# ════════════════════════════════════════════
# DOCX BUILD FONKSİYONLARI
# ════════════════════════════════════════════

def build_ba_docx(ba_content, project_name):
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)
    doc.add_heading(f'{project_name} — İş Analizi', level=0)
    doc.add_paragraph(f'Tarih: {time.strftime("%d.%m.%Y %H:%M")}')

    for i, e in enumerate(ba_content.get("ekranlar", []), 1):
        doc.add_heading(f'{i}. {e.get("ekran_adi", "")}', level=1)

        doc.add_heading(f'{i}.1. Açıklama', level=2)
        doc.add_paragraph(e.get("aciklama", ""))

        doc.add_heading(f'{i}.2. Gerekli Dokümanlar', level=2)
        doc.add_heading(f'{i}.2.1. Teknik Akış/Analiz', level=3)
        doc.add_paragraph(e.get("gerekli_dokumanlar", {}).get("teknik_akis", ""))
        doc.add_heading(f'{i}.2.2. Tasarım Dosyası', level=3)
        doc.add_paragraph(e.get("gerekli_dokumanlar", {}).get("tasarim_dosyasi", ""))

        doc.add_heading(f'{i}.3. İş Akışı Diyagramı', level=2)
        for adim in e.get("is_akisi_diyagrami", []):
            doc.add_paragraph(adim, style='List Number')

        frs = e.get("fonksiyonel_gereksinimler", [])
        doc.add_heading(f'{i}.4. Fonksiyonel Gereksinimler', level=2)
        for fr in frs:
            p = doc.add_paragraph()
            run = p.add_run(f'{fr.get("id","")}: ')
            run.bold = True
            p.add_run(fr.get("tanim", ""))

        kurallar = e.get("is_kurallari", [])
        doc.add_heading(f'{i}.5. İş Kuralları', level=2)
        for k in kurallar:
            p = doc.add_paragraph()
            run = p.add_run(k.get("kural", ""))
            run.bold = True
            if k.get("detay"):
                doc.add_paragraph(k["detay"])

        brs = e.get("kabul_kriterleri", [])
        doc.add_heading(f'{i}.6. Kabul Kriterleri', level=2)
        for br in brs:
            p = doc.add_paragraph()
            run = p.add_run(f'{br.get("id","")}: ')
            run.bold = True
            p.add_run(br.get("kriter", ""))

        vals = e.get("validasyonlar", [])
        doc.add_heading(f'{i}.7. Validasyonlar', level=2)
        for v in vals:
            doc.add_paragraph(
                f'{v.get("alan","")}: {v.get("kisit","")} → {v.get("hata_mesaji","")}',
                style='List Bullet'
            )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_ta_docx(ta_content, project_name):
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)

    ta = ta_content.get("teknik_analiz", ta_content)
    doc.add_heading(f'{project_name} — Teknik Analiz', level=0)
    doc.add_paragraph(f'Oluşturma Tarihi: {time.strftime("%d.%m.%Y %H:%M")}')

    genel = ta.get("genel_tanim", {})
    if genel:
        doc.add_heading("Genel Tanım", level=1)
        doc.add_paragraph(f'Modül: {genel.get("modul_adi", "")}')
        doc.add_paragraph(f'Stack: {", ".join(genel.get("teknoloji_stack", []))}')
        doc.add_paragraph(f'Mimari: {genel.get("mimari_yaklasim", "")}')

    endpoints = ta.get("endpoint_detaylari", {})
    if endpoints:
        doc.add_heading("API Endpoint Detayları", level=1)
        for ep, detail in endpoints.items():
            doc.add_heading(f'{detail.get("method", "GET")} {ep}', level=2)
            doc.add_paragraph(detail.get("aciklama", ""))
            if detail.get("request_body"):
                doc.add_paragraph(f'Request: {json.dumps(detail["request_body"], ensure_ascii=False, indent=2)[:500]}')
            if detail.get("response_success"):
                doc.add_paragraph(f'Response: {json.dumps(detail["response_success"], ensure_ascii=False, indent=2)[:500]}')
            for err in detail.get("response_errors", []):
                doc.add_paragraph(f'{err.get("http_code", "")} — {err.get("error_code", "")}: {err.get("mesaj", "")}', style='List Bullet')

    dtos = ta.get("dto_veri_yapilari", [])
    if dtos:
        doc.add_heading("DTO Veri Yapıları", level=1)
        for dto in dtos:
            doc.add_heading(dto.get("dto_adi", ""), level=2)
            doc.add_paragraph(dto.get("aciklama", ""))
            for f in dto.get("fields", []):
                doc.add_paragraph(f'{f.get("field", "")}: {f.get("tip", "")} {"(zorunlu)" if f.get("zorunlu") else ""} — {f.get("validasyon", "")}', style='List Bullet')

    vrs = ta.get("validasyon_kurallari", [])
    if vrs:
        doc.add_heading("Validasyon Kuralları", level=1)
        for v in vrs:
            doc.add_paragraph(f'{v.get("id", "")} [{v.get("field", "")}]: {v.get("kural", "")} → {v.get("hata_mesaji", "")}', style='List Bullet')

    curls = ta.get("mock_curl_ornekleri", [])
    if curls:
        doc.add_heading("cURL Örnekleri", level=1)
        for c in curls:
            doc.add_heading(c.get("endpoint_adi", ""), level=2)
            doc.add_paragraph(c.get("curl", ""))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_tc_excel(tc_content):
    """TC'yi 23-kolon Loodos şablonunda Excel'e çevir."""
    from utils.export import export_tc_excel
    return export_tc_excel(tc_content)


def build_tc_csv(tc_content):
    """TC'yi CSV'ye çevir."""
    tcs = tc_content.get("test_cases", [])
    if not tcs:
        return ""
    cols_23 = [
        "existance", "created_by", "date", "app_bundle", "test_case_id",
        "br_id", "tr_id", "priority", "channel", "testcase_type",
        "user_type", "test_area", "test_scenario", "testcase", "test_steps",
        "precondition", "test_data", "expected_result", "postcondition",
        "actual_result", "status", "regression_case", "comments"
    ]
    df = pd.DataFrame(tcs)
    for c in cols_23:
        if c not in df.columns:
            df[c] = ""
    df = df[cols_23]
    return df.to_csv(index=False, encoding="utf-8-sig")


def render_pipeline_result(ba_content, ta_content, tc_content, ba_qa, ta_qa, tc_qa,
                           ba_score, ta_score, tc_score, project_name, result_index, openapi_spec=None):
    """Tek bir pipeline sonucunu render et."""

    # SKOR KARTLARI
    col1, col2, col3 = st.columns(3)
    with col1:
        color = "🟢" if ba_score >= 60 else "🟡" if ba_score >= 40 else "🔴"
        st.metric(f"{color} BA Skor", f"{ba_score:.0f}/100")
    with col2:
        color = "🟢" if ta_score >= 60 else "🟡" if ta_score >= 40 else "🔴"
        st.metric(f"{color} TA Skor", f"{ta_score:.0f}/100")
    with col3:
        tc_count = len(tc_content.get("test_cases", [])) if tc_content else 0
        color = "🟢" if tc_score >= 60 else "🟡" if tc_score >= 40 else "🔴"
        st.metric(f"{color} TC Skor ({tc_count} TC)", f"{tc_score:.0f}/100")

    st.divider()

    # TABS
    tab_ba, tab_ta, tab_tc, tab_qa, tab_export = st.tabs(
        ["📋 İş Analizi", "⚙️ Teknik Analiz", "🧪 Test Case", "🔍 QA Detay", "📥 Export"]
    )

    # ── BA TAB ──
    with tab_ba:
        if ba_content:
            ekranlar = ba_content.get("ekranlar", [])
            st.subheader(f"İş Analizi — {len(ekranlar)} Ekran")
            for i, e in enumerate(ekranlar, 1):
                with st.expander(f"**{i}. {e.get('ekran_adi', '')}**", expanded=i == 1):
                    st.markdown(e.get("aciklama", ""))
                    akis = e.get("is_akisi_diyagrami", [])
                    if akis:
                        st.markdown("**İş Akışı:**")
                        for a in akis:
                            st.markdown(f"- {a}")
                    frs = e.get("fonksiyonel_gereksinimler", [])
                    if frs:
                        st.markdown(f"**Fonksiyonel Gereksinimler ({len(frs)}):**")
                        for fr in frs:
                            st.markdown(f"- **{fr.get('id','')}** {fr.get('tanim','')}")
                    kurallar = e.get("is_kurallari", [])
                    if kurallar:
                        st.markdown(f"**İş Kuralları ({len(kurallar)}):**")
                        for k in kurallar:
                            st.markdown(f"- {k.get('kural','')}")
                    brs = e.get("kabul_kriterleri", [])
                    if brs:
                        st.markdown(f"**Kabul Kriterleri ({len(brs)}):**")
                        for br in brs:
                            st.markdown(f"- **{br.get('id','')}** {br.get('kriter','')}")
                    vals = e.get("validasyonlar", [])
                    if vals:
                        st.markdown(f"**Validasyonlar ({len(vals)}):**")
                        for v in vals:
                            st.markdown(f"- `{v.get('alan','')}`: {v.get('kisit','')} → {v.get('hata_mesaji','')}")
        else:
            st.info("BA verisi mevcut değil.")

    # ── TA TAB ──
    with tab_ta:
        if ta_content:
            ta = ta_content.get("teknik_analiz", ta_content)
            st.subheader("Teknik Analiz")
            if ta.get("genel_tanim"):
                g = ta["genel_tanim"]
                st.markdown(f"**Modül:** {g.get('modul_adi', '')} | **Stack:** {', '.join(g.get('teknoloji_stack', []))}")
            if ta.get("endpoint_detaylari"):
                st.markdown(f"### API Endpoint'ler ({len(ta['endpoint_detaylari'])})")
                for ep, detail in ta["endpoint_detaylari"].items():
                    with st.expander(f"`{detail.get('method', 'GET')} {ep}`"):
                        st.markdown(detail.get("aciklama", ""))
                        if detail.get("request_body"):
                            st.json(detail["request_body"])
                        if detail.get("response_success"):
                            st.json(detail["response_success"])
            if ta.get("validasyon_kurallari"):
                st.markdown(f"### Validasyon Kuralları ({len(ta['validasyon_kurallari'])})")
                for v in ta["validasyon_kurallari"]:
                    st.markdown(f"- **{v.get('id', '')}** `{v.get('field', '')}`: {v.get('kural', '')}")
        else:
            st.info("TA verisi mevcut değil.")

    # ── TC TAB ──
    with tab_tc:
        if tc_content:
            test_cases = tc_content.get("test_cases", [])
            st.subheader(f"Test Cases — {len(test_cases)} adet")
            if test_cases:
                df = pd.DataFrame(test_cases)
                display_cols = [c for c in ["test_case_id", "priority", "test_area", "testcase", "test_steps", "expected_result"] if c in df.columns]
                st.dataframe(df[display_cols] if display_cols else df, use_container_width=True, height=400)
        else:
            st.info("TC verisi mevcut değil.")

    # ── QA TAB ──
    with tab_qa:
        st.subheader("🔍 QA Hakem Değerlendirmeleri")

        # Tab'lar ile ayrı ayrı göster
        if ba_qa or ta_qa or tc_qa:
            qa_tab_ba, qa_tab_ta, qa_tab_tc = st.tabs(["📋 BA Hakem", "⚙️ TA Hakem", "🧪 TC Hakem"])

            with qa_tab_ba:
                if ba_qa:
                    score = ba_qa.get("genel_puan", 0)
                    passed = ba_qa.get("gecti_mi", False)

                    # Genel değerlendirme
                    if passed:
                        st.success(f"✅ BA QA Geçti — Puan: {score}/100")
                    else:
                        st.warning(f"⚠️ BA QA Geçmedi — Puan: {score}/100")

                    # Genel değerlendirme metni
                    if ba_qa.get("genel_degerlendirme"):
                        st.markdown(f"**Genel Değerlendirme:** {ba_qa.get('genel_degerlendirme')}")

                    st.markdown("---")
                    st.markdown("**Kriter Skorları:**")

                    for s in ba_qa.get("skorlar", []):
                        puan = s.get("puan", 0)
                        bar = "🟩" * puan + "⬜" * (10 - puan)
                        st.markdown(f"**{s.get('kriter', '')}**: {bar} {puan}/10")
                        aciklama = s.get("aciklama", "")
                        # Fix: Convert dict/list to string
                        if isinstance(aciklama, dict):
                            aciklama = str(aciklama)
                        elif isinstance(aciklama, list):
                            aciklama = ', '.join(str(x) for x in aciklama)
                        if aciklama:
                            st.caption(aciklama)

                    oneriler = ba_qa.get("iyilestirme_onerileri", [])
                    if oneriler:
                        st.markdown("---")
                        st.markdown("**İyileştirme Önerileri:**")
                        for o in oneriler:
                            st.markdown(f"- {o}")
                else:
                    st.info("BA hakem değerlendirmesi mevcut değil.")

            with qa_tab_ta:
                if ta_qa:
                    score = ta_qa.get("genel_puan", 0)
                    passed = ta_qa.get("gecti_mi", False)

                    # Genel değerlendirme
                    if passed:
                        st.success(f"✅ TA QA Geçti — Puan: {score}/100")
                    else:
                        st.warning(f"⚠️ TA QA Geçmedi — Puan: {score}/100")

                    # Genel değerlendirme metni
                    if ta_qa.get("genel_degerlendirme"):
                        st.markdown(f"**Genel Değerlendirme:** {ta_qa.get('genel_degerlendirme')}")

                    st.markdown("---")
                    st.markdown("**Kriter Skorları:**")

                    for s in ta_qa.get("skorlar", []):
                        puan = s.get("puan", 0)
                        bar = "🟩" * puan + "⬜" * (10 - puan)
                        st.markdown(f"**{s.get('kriter', '')}**: {bar} {puan}/10")
                        aciklama = s.get("aciklama", "")
                        # Fix: Convert dict/list to string
                        if isinstance(aciklama, dict):
                            aciklama = str(aciklama)
                        elif isinstance(aciklama, list):
                            aciklama = ', '.join(str(x) for x in aciklama)
                        if aciklama:
                            st.caption(aciklama)

                    oneriler = ta_qa.get("iyilestirme_onerileri", [])
                    if oneriler:
                        st.markdown("---")
                        st.markdown("**İyileştirme Önerileri:**")
                        for o in oneriler:
                            st.markdown(f"- {o}")
                else:
                    st.info("TA hakem değerlendirmesi mevcut değil.")

            with qa_tab_tc:
                if tc_qa:
                    score = tc_qa.get("genel_puan", 0)
                    passed = tc_qa.get("gecti_mi", False)

                    # Genel değerlendirme
                    if passed:
                        st.success(f"✅ TC QA Geçti — Puan: {score}/100")
                    else:
                        st.warning(f"⚠️ TC QA Geçmedi — Puan: {score}/100")

                    # Genel değerlendirme metni
                    if tc_qa.get("genel_degerlendirme"):
                        st.markdown(f"**Genel Değerlendirme:** {tc_qa.get('genel_degerlendirme')}")

                    st.markdown("---")
                    st.markdown("**Kriter Skorları:**")

                    for s in tc_qa.get("skorlar", []):
                        puan = s.get("puan", 0)
                        bar = "🟩" * puan + "⬜" * (10 - puan)
                        st.markdown(f"**{s.get('kriter', '')}**: {bar} {puan}/10")
                        aciklama = s.get("aciklama", "")
                        # Fix: Convert dict/list to string
                        if isinstance(aciklama, dict):
                            aciklama = str(aciklama)
                        elif isinstance(aciklama, list):
                            aciklama = ', '.join(str(x) for x in aciklama)
                        if aciklama:
                            st.caption(aciklama)

                    oneriler = tc_qa.get("iyilestirme_onerileri", [])
                    if oneriler:
                        st.markdown("---")
                        st.markdown("**İyileştirme Önerileri:**")
                        for o in oneriler:
                            st.markdown(f"- {o}")
                else:
                    st.info("TC hakem değerlendirmesi mevcut değil.")
        else:
            st.info("QA hakem değerlendirmesi mevcut değil.")

    # ── EXPORT TAB ──
    with tab_export:
        st.subheader("📥 Dokümanları İndir")

        col1, col2, col3 = st.columns(3)

        # BA → DOCX
        with col1:
            st.markdown("### 📋 İş Analizi")
            if ba_content:
                ba_docx = build_ba_docx(ba_content, project_name)
                st.download_button(
                    "📥 İş Analizi (.docx)",
                    data=ba_docx,
                    file_name=f"{project_name}_is_analizi.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key=f"ba_download_{result_index}"
                )
            else:
                st.info("BA verisi mevcut değil.")

        # TA → DOCX + OpenAPI
        with col2:
            st.markdown("### ⚙️ Teknik Analiz")
            if ta_content:
                ta_docx = build_ta_docx(ta_content, project_name)
                st.download_button(
                    "📥 Teknik Analiz (.docx)",
                    data=ta_docx,
                    file_name=f"{project_name}_teknik_analiz.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key=f"ta_download_{result_index}"
                )

                # OpenAPI Spec export (if available)
                if openapi_spec:
                    openapi_json = json.dumps(openapi_spec, ensure_ascii=False, indent=2) if isinstance(openapi_spec, dict) else openapi_spec
                    st.download_button(
                        "📥 OpenAPI Spec (.json)",
                        data=openapi_json,
                        file_name=f"{project_name}_openapi.json",
                        mime="application/json",
                        use_container_width=True,
                        key=f"openapi_download_{result_index}"
                    )
            else:
                st.info("TA verisi mevcut değil.")

        # TC → CSV + Excel
        with col3:
            st.markdown("### 🧪 Test Cases")
            if tc_content and tc_content.get("test_cases"):
                tc_csv = build_tc_csv(tc_content)
                st.download_button(
                    "📥 Test Cases (.csv)",
                    data=tc_csv,
                    file_name=f"{project_name}_test_cases.csv",
                    mime="text/csv",
                    use_container_width=True,
                    key=f"tc_csv_download_{result_index}"
                )
                tc_xlsx = build_tc_excel(tc_content)
                st.download_button(
                    "📥 Test Cases (.xlsx) — 23 kolon",
                    data=tc_xlsx,
                    file_name=f"{project_name}_test_cases.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                    key=f"tc_xlsx_download_{result_index}"
                )
            else:
                st.info("TC verisi mevcut değil.")


# ════════════════════════════════════════════
# ANA SAYFA
# ════════════════════════════════════════════

# Önce session state'ten son çalıştırmayı kontrol et
current_run = None
ba_content = st.session_state.get("ba_content")
ta_content = st.session_state.get("ta_content")
tc_content = st.session_state.get("tc_content")
ba_qa = st.session_state.get("ba_qa_result", {})
ta_qa = st.session_state.get("ta_qa_result", {})
tc_qa = st.session_state.get("tc_qa_result", {})
ba_score = st.session_state.get("ba_score", 0)
ta_score = st.session_state.get("ta_score", 0)
tc_score = st.session_state.get("tc_score", 0)
project = st.session_state.get("project_name", "Proje")

# Fallback: last_pipeline_result
if not ba_content:
    results = st.session_state.get("last_pipeline_result")
    if results:
        ba_content = results.get("ba", {}).get("content")
        ta_content = results.get("ta", {}).get("content")
        tc_content = results.get("tc", {}).get("content")
        ba_qa = results.get("ba", {}).get("qa_result", {})
        ta_qa = results.get("ta", {}).get("qa_result", {})
        tc_qa = results.get("tc", {}).get("qa_result", {})
        ba_score = results.get("ba", {}).get("quality_score", 0)
        ta_score = results.get("ta", {}).get("quality_score", 0)
        tc_score = results.get("tc", {}).get("quality_score", 0)
        project = results.get("project_name", "Proje")

# Session state'te aktif çalışma varsa onu göster
if ba_content:
    current_run = {
        "ba_content": ba_content,
        "ta_content": ta_content,
        "tc_content": tc_content,
        "ba_qa": ba_qa,
        "ta_qa": ta_qa,
        "tc_qa": tc_qa,
        "ba_score": ba_score,
        "ta_score": ta_score,
        "tc_score": tc_score,
        "project_name": project,
    }

# Database'den geçmiş çalıştırmaları çek
try:
    from data.database import get_recent_pipeline_runs as get_recent_runs, get_pipeline_run_outputs as get_run_outputs

    past_runs = get_recent_runs(50)  # Son 50 çalıştırma

    if not past_runs and not current_run:
        st.info("Henüz pipeline çalıştırılmadı veya tamamlanmadı. Önce **BRD Pipeline** sayfasından çalıştırın.")
        st.stop()

    # İstatistikler
    total_runs = len(past_runs)
    if current_run:
        total_runs += 1

    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Toplam Çalıştırma", total_runs)
    with col2:
        if past_runs:
            avg_ba = sum(r["ba_score"] for r in past_runs) / len(past_runs)
            st.metric("Ortalama BA Skoru", f"{avg_ba:.0f}")
    with col3:
        if past_runs:
            completed = sum(1 for r in past_runs if r["status"] == "completed")
            st.metric("Tamamlanan", f"{completed}/{len(past_runs)}")

    st.divider()

    # Son çalıştırma (session state'ten)
    if current_run:
        st.subheader("🔥 Son Çalıştırma")

        # OpenAPI spec'i session state'ten veya database'den al
        openapi_spec = None
        if st.session_state.get("openapi_spec"):
            openapi_spec = st.session_state.get("openapi_spec")
        elif st.session_state.get("openapi_json"):
            try:
                openapi_spec = json.loads(st.session_state.get("openapi_json"))
            except:
                openapi_spec = None

        with st.expander(f"**{current_run['project_name']}** — BA:{current_run['ba_score']:.0f} | TA:{current_run['ta_score']:.0f} | TC:{current_run['tc_score']:.0f}", expanded=True):
            render_pipeline_result(
                current_run["ba_content"],
                current_run["ta_content"],
                current_run["tc_content"],
                current_run["ba_qa"],
                current_run["ta_qa"],
                current_run["tc_qa"],
                current_run["ba_score"],
                current_run["ta_score"],
                current_run["tc_score"],
                current_run["project_name"],
                "current",
                openapi_spec
            )
        st.divider()

    # Geçmiş çalıştırmalar
    if past_runs:
        st.subheader("📜 Geçmiş Çalıştırmalar")

        for idx, run in enumerate(past_runs):
            status_icon = "✅" if run["status"] == "completed" else "🔄" if run["status"] == "running" else "❌"
            avg_score = (run["ba_score"] + run["ta_score"] + run["tc_score"]) / 3 if run["ba_score"] else 0

            with st.expander(
                f"{status_icon} **{run['project_name']}** — "
                f"BA:{run['ba_score']:.0f} | TA:{run['ta_score']:.0f} | TC:{run['tc_score']:.0f} | "
                f"Ort:{avg_score:.0f} — {run['created_at'][:16]}"
            ):
                # Pipeline detaylarını göster
                col1, col2, col3, col4, col5 = st.columns(5)
                col1.metric("BA Skor", f"{run['ba_score']:.0f}", delta=f"{run['ba_revisions']} rev")
                col2.metric("TA Skor", f"{run['ta_score']:.0f}", delta=f"{run['ta_revisions']} rev")
                col3.metric("TC Skor", f"{run['tc_score']:.0f}", delta=f"{run['tc_revisions']} rev")
                col4.metric("Durum", run["status"])
                col5.metric("Süre", f"{run['total_time_sec']}s")

                # Stage output'larını al
                outputs = get_run_outputs(run["id"])

                # En son BA, TA, TC çıktılarını bul
                ba_output = next((o for o in reversed(outputs) if o["stage"] == "ba"), None)
                ta_output = next((o for o in reversed(outputs) if o["stage"] == "ta"), None)
                tc_output = next((o for o in reversed(outputs) if o["stage"] == "tc"), None)

                # İçerikleri parse et
                ba_content_db = json.loads(ba_output["content_json"]) if ba_output and ba_output["content_json"] else None
                ta_content_db = json.loads(ta_output["content_json"]) if ta_output and ta_output["content_json"] else None
                tc_content_db = json.loads(tc_output["content_json"]) if tc_output and tc_output["content_json"] else None

                ba_qa_db = json.loads(ba_output["qa_result_json"]) if ba_output and ba_output["qa_result_json"] else {}
                ta_qa_db = json.loads(ta_output["qa_result_json"]) if ta_output and ta_output["qa_result_json"] else {}
                tc_qa_db = json.loads(tc_output["qa_result_json"]) if tc_output and tc_output["qa_result_json"] else {}

                # OpenAPI spec'i parse et
                openapi_spec_db = None
                if ta_output and ta_output.get("openapi_spec_json"):
                    try:
                        openapi_spec_db = json.loads(ta_output["openapi_spec_json"])
                    except:
                        openapi_spec_db = None

                # Render et
                if ba_content_db or ta_content_db or tc_content_db:
                    render_pipeline_result(
                        ba_content_db,
                        ta_content_db,
                        tc_content_db,
                        ba_qa_db,
                        ta_qa_db,
                        tc_qa_db,
                        run["ba_score"],
                        run["ta_score"],
                        run["tc_score"],
                        run["project_name"],
                        f"past_{idx}",
                        openapi_spec_db
                    )
                else:
                    st.warning("Bu pipeline çalıştırmasına ait çıktı bulunamadı.")

except ImportError:
    # Database yok, sadece session state'i göster
    if current_run:
        # OpenAPI spec'i session state'ten al
        openapi_spec = None
        if st.session_state.get("openapi_spec"):
            openapi_spec = st.session_state.get("openapi_spec")
        elif st.session_state.get("openapi_json"):
            try:
                openapi_spec = json.loads(st.session_state.get("openapi_json"))
            except:
                openapi_spec = None

        render_pipeline_result(
            current_run["ba_content"],
            current_run["ta_content"],
            current_run["tc_content"],
            current_run["ba_qa"],
            current_run["ta_qa"],
            current_run["tc_qa"],
            current_run["ba_score"],
            current_run["ta_score"],
            current_run["tc_score"],
            current_run["project_name"],
            "current",
            openapi_spec
        )
    else:
        st.info("Henüz pipeline çalıştırılmadı veya tamamlanmadı. Önce **BRD Pipeline** sayfasından çalıştırın.")
