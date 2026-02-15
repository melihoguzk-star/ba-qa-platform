"""Sayfa 1: BRD Pipeline — Adım adım, kullanıcı onaylı akış.

Tüm step'ler elif zincirinde — her render'da SADECE 1 step çalışır.
DOCX build fonksiyonları dosya içinde tanımlı (utils.export'a bağımlılık yok).
"""
import json
import time
import io
import streamlit as st
from utils.config import (BA_PASS_THRESHOLD, TA_PASS_THRESHOLD, TC_PASS_THRESHOLD, MAX_REVISIONS,
                          ANTHROPIC_MODELS, GEMINI_MODELS, ALL_MODELS, SONNET_MODEL, GEMINI_MODEL)
from components.sidebar import render_custom_sidebar

st.set_page_config(page_title="BRD Pipeline", page_icon="📋", layout="wide")
render_custom_sidebar(active_page="brd_pipeline")
st.title("📋 BRD Pipeline — Adım Adım")

# ── API Key kontrolü ──
anthropic_key = st.session_state.get("anthropic_api_key", "")
gemini_key = st.session_state.get("gemini_key", "")

# Session state'de yoksa secrets'tan dene
if not anthropic_key:
    try:
        anthropic_key = st.secrets.get("ANTHROPIC_API_KEY", "")
    except Exception:
        pass

if not gemini_key:
    try:
        gemini_key = st.secrets.get("GEMINI_API_KEY", "")
    except Exception:
        pass

# Hala yoksa hata ver
if not anthropic_key or not gemini_key:
    st.error("⚠️ API Key'ler bulunamadı!")
    st.info("🔑 Lütfen `.streamlit/secrets.toml` dosyasına ANTHROPIC_API_KEY ve GEMINI_API_KEY ekleyin.")
    st.code("""# .streamlit/secrets.toml
ANTHROPIC_API_KEY = "your-key-here"
GEMINI_API_KEY = "your-key-here"
""", language="toml")
    st.stop()

# Session state'e yaz (diğer sayfalar için)
st.session_state["anthropic_api_key"] = anthropic_key
st.session_state["gemini_key"] = gemini_key

if "pipeline_step" not in st.session_state:
    st.session_state.pipeline_step = "upload"


# ════════════════════════════════════════════
# LOG HELPERS
# ════════════════════════════════════════════
def get_log():
    if "log_messages" not in st.session_state:
        st.session_state.log_messages = []
    return st.session_state.log_messages

def log(msg):
    get_log().append(msg)

def show_log():
    msgs = get_log()
    if msgs:
        with st.expander("📜 Pipeline Log", expanded=False):
            for m in msgs:
                st.text(m)

STEP_ORDER = ["upload", "ba_gen", "ba_review", "ba_qa", "ta_gen", "ta_review", "ta_qa", "tc_gen", "tc_review", "tc_qa", "done"]
def show_progress():
    step = st.session_state.pipeline_step
    if step in STEP_ORDER:
        idx = STEP_ORDER.index(step)
        st.progress(idx / (len(STEP_ORDER) - 1), text=f"Adım {idx}/{len(STEP_ORDER) - 1}")


# ════════════════════════════════════════════
# DOCX BUILDERS (inline — utils.export'a bağımlılık yok)
# ════════════════════════════════════════════
def build_ba_docx(ba_content, project_name):
    from docx import Document as DocxDocument
    from docx.shared import Pt
    doc = DocxDocument()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)
    doc.add_heading(f'{project_name} — İş Analizi', level=0)
    doc.add_paragraph(f'Tarih: {time.strftime("%d.%m.%Y %H:%M")}')

    for i, e in enumerate(ba_content.get("ekranlar", []), 1):
        # 1. Geliştirmenin Adı
        doc.add_heading(f'{i}. {e.get("ekran_adi", "")}', level=1)

        # 1.1. Açıklama
        doc.add_heading(f'{i}.1. Açıklama', level=2)
        doc.add_paragraph(e.get("aciklama", ""))

        # 1.2. Gerekli Dokümanlar (boş)
        doc.add_heading(f'{i}.2. Gerekli Dokümanlar', level=2)
        doc.add_heading(f'{i}.2.1. Teknik Akış/Analiz', level=3)
        doc.add_paragraph(e.get("gerekli_dokumanlar", {}).get("teknik_akis", ""))
        doc.add_heading(f'{i}.2.2. Tasarım Dosyası', level=3)
        doc.add_paragraph(e.get("gerekli_dokumanlar", {}).get("tasarim_dosyasi", ""))

        # 1.3. İş Akışı Diyagramı
        doc.add_heading(f'{i}.3. İş Akışı Diyagramı', level=2)
        for adim in e.get("is_akisi_diyagrami", []):
            doc.add_paragraph(adim, style='List Number')

        # 1.4. Fonksiyonel Gereksinimler
        frs = e.get("fonksiyonel_gereksinimler", [])
        doc.add_heading(f'{i}.4. Fonksiyonel Gereksinimler', level=2)
        for fr in frs:
            p = doc.add_paragraph()
            run = p.add_run(f'{fr.get("id","")}: ')
            run.bold = True
            p.add_run(fr.get("tanim", ""))

        # 1.5. İş Kuralları
        kurallar = e.get("is_kurallari", [])
        doc.add_heading(f'{i}.5. İş Kuralları', level=2)
        for k in kurallar:
            p = doc.add_paragraph()
            run = p.add_run(k.get("kural", ""))
            run.bold = True
            if k.get("detay"):
                doc.add_paragraph(k["detay"])

        # 1.6. Kabul Kriterleri
        brs = e.get("kabul_kriterleri", [])
        doc.add_heading(f'{i}.6. Kabul Kriterleri', level=2)
        for br in brs:
            p = doc.add_paragraph()
            run = p.add_run(f'{br.get("id","")}: ')
            run.bold = True
            p.add_run(br.get("kriter", ""))

        # 1.7. Validasyonlar
        vals = e.get("validasyonlar", [])
        doc.add_heading(f'{i}.7. Validasyonlar', level=2)
        for v in vals:
            doc.add_paragraph(
                f'{v.get("alan","")}: {v.get("kisit","")} → {v.get("hata_mesaji","")}',
                style='List Bullet'
            )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_ta_docx(ta_content, project_name):
    from docx import Document as DocxDocument
    from docx.shared import Pt
    doc = DocxDocument()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)
    ta = ta_content.get("teknik_analiz", ta_content)
    doc.add_heading(f'{project_name} — Teknik Analiz', level=0)
    doc.add_paragraph(f'Tarih: {time.strftime("%d.%m.%Y %H:%M")}')
    genel = ta.get("genel_tanim", {})
    if genel:
        doc.add_heading("Genel Tanım", level=1)
        doc.add_paragraph(f'Modül: {genel.get("modul_adi","")}')
        doc.add_paragraph(f'Stack: {", ".join(genel.get("teknoloji_stack",[]))}')
        doc.add_paragraph(f'Mimari: {genel.get("mimari_yaklasim","")}')
    endpoints = ta.get("endpoint_detaylari", {})
    if endpoints:
        doc.add_heading("API Endpoints", level=1)
        for ep, d in endpoints.items():
            doc.add_heading(f'{d.get("method","GET")} {ep}', level=2)
            doc.add_paragraph(d.get("aciklama", ""))
            if d.get("request_body"):
                doc.add_paragraph(f'Request: {json.dumps(d["request_body"], ensure_ascii=False, indent=2)[:500]}')
            if d.get("response_success"):
                doc.add_paragraph(f'Response: {json.dumps(d["response_success"], ensure_ascii=False, indent=2)[:500]}')
            for err in d.get("response_errors", []):
                doc.add_paragraph(f'{err.get("http_code","")} {err.get("error_code","")}: {err.get("mesaj","")}', style='List Bullet')
    dtos = ta.get("dto_veri_yapilari", [])
    if dtos:
        doc.add_heading("DTO Yapıları", level=1)
        for dto in dtos:
            doc.add_heading(dto.get("dto_adi",""), level=2)
            doc.add_paragraph(dto.get("aciklama",""))
            for f in dto.get("fields",[]):
                doc.add_paragraph(f'{f.get("field","")}: {f.get("tip","")} {f.get("validasyon","")}', style='List Bullet')
    vrs = ta.get("validasyon_kurallari", [])
    if vrs:
        doc.add_heading("Validasyon Kuralları", level=1)
        for v in vrs:
            doc.add_paragraph(f'{v.get("id","")} [{v.get("field","")}]: {v.get("kural","")}', style='List Bullet')
    curls = ta.get("mock_curl_ornekleri", [])
    if curls:
        doc.add_heading("cURL Örnekleri", level=1)
        for c in curls:
            doc.add_paragraph(c.get("endpoint_adi",""), style='List Bullet')
            doc.add_paragraph(c.get("curl",""))
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_tc_docx(tc_content, project_name):
    from docx import Document as DocxDocument
    from docx.shared import Pt
    doc = DocxDocument()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(9)
    doc.add_heading(f'{project_name} — Test Cases', level=0)
    tcs = tc_content.get("test_cases", [])
    if not tcs:
        doc.add_paragraph("Test case bulunamadı.")
    else:
        cols = ["test_case_id","br_id","priority","test_area","testcase","test_steps","test_data","expected_result"]
        headers = ["TC ID","BR","Priority","Area","Test Case","Steps","Data","Expected"]
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = 'Table Grid'
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            for p in table.rows[0].cells[i].paragraphs:
                for r in p.runs:
                    r.bold = True
        for tc in tcs:
            row = table.add_row()
            for i, col in enumerate(cols):
                row.cells[i].text = str(tc.get(col,""))[:200]
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ════════════════════════════════════════════
# PREVIEW HELPERS
# ════════════════════════════════════════════
def show_ba_preview(ba_content):
    ekranlar = ba_content.get("ekranlar", [])
    st.info(f"📋 **İş Analizi** — {len(ekranlar)} ekran")
    for i, e in enumerate(ekranlar, 1):
        with st.expander(f"**{i}. {e.get('ekran_adi','')}**", expanded=i <= 2):
            st.markdown(e.get("aciklama","")[:500])
            frs = e.get("fonksiyonel_gereksinimler",[])
            if frs:
                st.markdown(f"**Fonksiyonel Gereksinimler ({len(frs)}):**")
                for fr in frs[:5]:
                    st.markdown(f"- **{fr.get('id','')}** {fr.get('tanim','')}")
            kurallar = e.get("is_kurallari",[])
            if kurallar:
                st.markdown(f"**İş Kuralları ({len(kurallar)}):**")
                for k in kurallar[:3]:
                    st.markdown(f"- {k.get('kural','')}")
            brs = e.get("kabul_kriterleri",[])
            if brs:
                st.markdown(f"**Kabul Kriterleri ({len(brs)}):**")
                for br in brs[:3]:
                    st.markdown(f"- **{br.get('id','')}** {br.get('kriter','')}")
            vals = e.get("validasyonlar",[])
            if vals:
                st.markdown(f"**Validasyonlar ({len(vals)}):**")
                for v in vals[:3]:
                    st.markdown(f"- `{v.get('alan','')}`: {v.get('kisit','')}")

def show_ta_preview(ta_content):
    ta = ta_content.get("teknik_analiz", ta_content)
    ep = len(ta.get("endpoint_detaylari",{}))
    vr = len(ta.get("validasyon_kurallari",[]))
    st.info(f"⚙️ **Teknik Analiz** — {ep} endpoint, {vr} validasyon")
    if ta.get("endpoint_detaylari"):
        with st.expander(f"**Endpoints ({ep})**", expanded=True):
            for e, d in list(ta["endpoint_detaylari"].items())[:5]:
                st.markdown(f"- `{d.get('method','GET')} {e}` — {d.get('aciklama','')[:80]}")

def show_tc_preview(tc_content):
    tcs = tc_content.get("test_cases",[])
    st.info(f"🧪 **Test Case** — {len(tcs)} adet")
    if tcs:
        import pandas as pd
        df = pd.DataFrame(tcs)
        cols = [c for c in ["test_case_id","priority","test_area","testcase","expected_result"] if c in df.columns]
        st.dataframe(df[cols] if cols else df.head(10), use_container_width=True, height=400)

def show_qa_result(qa_result, stage_name):
    score = qa_result.get("genel_puan",0)
    passed = qa_result.get("gecti_mi", False)
    if passed:
        st.success(f"✅ {stage_name} QA Geçti — Puan: {score}/100")
    else:
        st.warning(f"⚠️ {stage_name} QA Geçmedi — Puan: {score}/100")
    for s in qa_result.get("skorlar",[]):
        puan = s.get("puan",0)
        bar = "🟩" * puan + "⬜" * (10 - puan)
        aciklama = s.get('aciklama', '')
        # Eğer aciklama dict ise string'e çevir
        if isinstance(aciklama, dict):
            aciklama = str(aciklama)
        elif isinstance(aciklama, list):
            aciklama = ', '.join(str(x) for x in aciklama)
        st.markdown(f"**{s.get('kriter','')}**: {bar} {puan}/10 — {aciklama}")
    st.markdown(f"**Değerlendirme:** {qa_result.get('genel_degerlendirme','')}")
    oneriler = qa_result.get("iyilestirme_onerileri",[])
    if oneriler:
        st.markdown("**Öneriler:** " + " | ".join(oneriler))


# ════════════════════════════════════════════
# ANA STEP AKIŞI — elif ZİNCİRİ
# ════════════════════════════════════════════
step = st.session_state.pipeline_step
if step != "upload":
    show_progress()

# ──── UPLOAD ────
if step == "upload":
    with st.form("brd_form"):
        st.subheader("📄 BRD Dosyası Yükle")
        c1, c2 = st.columns(2)
        with c1:
            project_name = st.text_input("Proje Adı *", placeholder="SunPass")
            jira_key = st.text_input("Jira Key", placeholder="SunPass")
        with c2:
            priority = st.selectbox("Öncelik", ["Kritik","Yüksek","Orta","Düşük"], index=2)
            uploaded_file = st.file_uploader("BRD Dosyası *", type=["pdf","docx","txt"])

        st.divider()
        st.markdown("#### 🤖 Model Seçimi")
        c3, c4 = st.columns(2)
        with c3:
            # Get default model index
            default_gen_idx = list(ALL_MODELS.values()).index(SONNET_MODEL) if SONNET_MODEL in ALL_MODELS.values() else 2
            generation_model = st.selectbox(
                "Generation Model (BA/TA/TC)",
                options=list(ALL_MODELS.keys()),
                index=default_gen_idx,
                help="Doküman üretimi için kullanılacak model (Anthropic veya Gemini)"
            )
        with c4:
            # Get default model index
            default_qa_idx = list(ALL_MODELS.values()).index(GEMINI_MODEL) if GEMINI_MODEL in ALL_MODELS.values() else 4
            qa_model = st.selectbox(
                "QA/Hakem Model",
                options=list(ALL_MODELS.keys()),
                index=default_qa_idx,
                help="Kalite değerlendirmesi için kullanılacak model (Anthropic veya Gemini)"
            )

        use_cp = st.checkbox("Checkpoint kullan (24h cache)", value=True)
        submitted = st.form_submit_button("🚀 Başlat", type="primary", use_container_width=True)
    if submitted:
        if not project_name or not uploaded_file:
            st.error("Proje adı ve BRD dosyası zorunlu.")
            st.stop()
        from utils.text_extractor import extract_text
        with st.spinner("Text çıkarılıyor..."):
            brd_text = extract_text(uploaded_file)
        if len(brd_text.strip()) < 100:
            st.error("BRD çok kısa veya boş.")
            st.stop()
        if not use_cp:
            from pipeline.brd.checkpoint import clear_all_checkpoints
            clear_all_checkpoints(project_name)
        from pipeline.brd.orchestrator import init_run
        st.session_state.brd_text = brd_text
        st.session_state.project_name = project_name
        st.session_state.jira_key = jira_key or project_name
        st.session_state.priority = priority
        st.session_state.generation_model = ALL_MODELS[generation_model]
        st.session_state.qa_model = ALL_MODELS[qa_model]
        st.session_state.run_id = init_run(project_name, jira_key or project_name, priority, uploaded_file.name)
        st.session_state.pipeline_start = time.time()
        st.session_state.pipeline_step = "ba_gen"
        st.session_state.log_messages = []
        st.rerun()

# ──── BA GENERATE ────
elif step == "ba_gen":
    st.subheader("📋 Adım 1/3 — İş Analizi Üretimi")
    from pipeline.brd.orchestrator import generate_ba
    gen_model = st.session_state.get("generation_model")
    with st.status("🤖 BA üretiliyor...", expanded=True) as s:
        ba = generate_ba(st.session_state.brd_text, st.session_state.project_name, anthropic_key, gemini_key, log,
                        st.session_state.get("ba_feedback",""), model=gen_model)
        s.update(label="✅ BA hazır!", state="complete")
    st.session_state.ba_content = ba
    st.session_state.pipeline_step = "ba_review"
    st.rerun()

# ──── BA REVIEW ────
elif step == "ba_review":
    st.subheader("📋 Adım 1/3 — İş Analizi İnceleme")
    show_ba_preview(st.session_state.ba_content)
    show_log()
    st.divider()

    # Hakeme gönderme seçeneği
    skip_qa = st.checkbox("⚡ Hakeme göndermeden devam et (QA'yı atla)", key="skip_ba_qa")

    c1, c2 = st.columns(2)
    with c1:
        if st.button("✅ Onayla ve İlerle", type="primary", use_container_width=True):
            if skip_qa:
                # QA'yı atla, direkt TA'ya geç - Database'e kaydet
                from pipeline.brd.orchestrator import finalize_stage
                qa_result = {"genel_puan": 100, "gecti_mi": True, "genel_degerlendirme": "QA atlandı (Force Pass)",
                           "skorlar": [], "iyilestirme_onerileri": []}
                finalize_stage(st.session_state.run_id, "ba", st.session_state.ba_content, qa_result,
                             st.session_state.get("ba_revision_count", 0), True, 0)
                st.session_state.ba_score = 100  # Force pass
                st.session_state.ba_qa_result = qa_result
                st.session_state.pipeline_step = "ta_gen"
                st.success("⚡ QA atlandı, Teknik Analize geçiliyor...")
            else:
                # Normal akış, QA'ya gönder
                st.session_state.pipeline_step = "ba_qa"
            st.rerun()
    with c2:
        if st.button("🔄 Yeniden Üret", use_container_width=True):
            from pipeline.brd.checkpoint import clear_checkpoint
            clear_checkpoint(st.session_state.project_name, "ba")
            st.session_state.pipeline_step = "ba_gen"
            st.rerun()

# ──── BA QA ────
elif step == "ba_qa":
    st.subheader("📋 Adım 1/3 — BA QA Hakem")
    if "ba_qa_result" not in st.session_state:
        from pipeline.brd.orchestrator import evaluate_ba_qa, finalize_stage
        from pipeline.brd.checkpoint import clear_checkpoint
        qa_model = st.session_state.get("qa_model")
        with st.status("🔍 BA QA değerlendiriyor...", expanded=True):
            qa = evaluate_ba_qa(st.session_state.ba_content, anthropic_key, gemini_key, log, model=qa_model)
        finalize_stage(st.session_state.run_id, "ba", st.session_state.ba_content, qa, st.session_state.get("ba_revision_count",0), False, 0)
        clear_checkpoint(st.session_state.project_name, "ba")
        st.session_state.ba_qa_result = qa
        st.session_state.ba_score = qa.get("genel_puan",0)

    qa = st.session_state.ba_qa_result
    score = st.session_state.ba_score
    show_qa_result(qa, "BA")
    show_log()

    # DOCX download
    st.divider()
    st.download_button("📥 BA Dokümanı İndir (.docx)",
        data=build_ba_docx(st.session_state.ba_content, st.session_state.project_name),
        file_name=f"{st.session_state.project_name}_is_analizi.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True)

    st.divider()
    if score >= BA_PASS_THRESHOLD:
        if st.button("➡️ Teknik Analize Geç", type="primary", use_container_width=True):
            st.session_state.pipeline_step = "ta_gen"
            st.rerun()
    else:
        rev = st.session_state.get("ba_revision_count",0)
        if rev < MAX_REVISIONS:
            c1, c2 = st.columns(2)
            with c1:
                if st.button("🔄 Revize Et", type="primary", use_container_width=True):
                    st.session_state.ba_revision_count = rev + 1
                    st.session_state.ba_feedback = qa.get("genel_degerlendirme","") + " | " + ", ".join(qa.get("iyilestirme_onerileri",[]))
                    del st.session_state["ba_qa_result"]
                    from pipeline.brd.checkpoint import clear_checkpoint
                    clear_checkpoint(st.session_state.project_name, "ba")
                    st.session_state.pipeline_step = "ba_gen"
                    st.rerun()
            with c2:
                if st.button("➡️ Yine de Devam Et", use_container_width=True):
                    st.session_state.pipeline_step = "ta_gen"
                    st.rerun()
        else:
            st.warning("Max revizyon.")
            if st.button("➡️ Teknik Analize Geç", type="primary", use_container_width=True):
                st.session_state.pipeline_step = "ta_gen"
                st.rerun()

# ──── TA GENERATE ────
elif step == "ta_gen":
    st.subheader("⚙️ Adım 2/3 — Teknik Analiz Üretimi")
    from pipeline.brd.orchestrator import generate_ta
    gen_model = st.session_state.get("generation_model")
    with st.status("🤖 TA üretiliyor...", expanded=True) as s:
        ta = generate_ta(st.session_state.brd_text, st.session_state.ba_content, st.session_state.project_name,
                        anthropic_key, gemini_key, log, st.session_state.get("ta_feedback",""), model=gen_model)
        s.update(label="✅ TA hazır!", state="complete")
    st.session_state.ta_content = ta
    st.session_state.pipeline_step = "ta_review"
    st.rerun()

# ──── TA REVIEW ────
elif step == "ta_review":
    st.subheader("⚙️ Adım 2/3 — Teknik Analiz İnceleme")
    show_ta_preview(st.session_state.ta_content)
    show_log()
    st.divider()

    # Hakeme gönderme seçeneği
    skip_qa = st.checkbox("⚡ Hakeme göndermeden devam et (QA'yı atla)", key="skip_ta_qa")

    c1, c2 = st.columns(2)
    with c1:
        if st.button("✅ Onayla ve İlerle", type="primary", use_container_width=True, key="ta_ok"):
            if skip_qa:
                # QA'yı atla, direkt TC'ye geç - Database'e kaydet
                from pipeline.brd.orchestrator import finalize_stage
                qa_result = {"genel_puan": 100, "gecti_mi": True, "genel_degerlendirme": "QA atlandı (Force Pass)",
                           "skorlar": [], "iyilestirme_onerileri": []}
                finalize_stage(st.session_state.run_id, "ta", st.session_state.ta_content, qa_result,
                             st.session_state.get("ta_revision_count", 0), True, 0)
                st.session_state.ta_score = 100  # Force pass
                st.session_state.ta_qa_result = qa_result
                st.session_state.pipeline_step = "tc_gen"
                st.success("⚡ QA atlandı, Test Case'e geçiliyor...")
            else:
                # Normal akış, QA'ya gönder
                st.session_state.pipeline_step = "ta_qa"
            st.rerun()
    with c2:
        if st.button("🔄 Yeniden Üret", use_container_width=True, key="ta_re"):
            from pipeline.brd.checkpoint import clear_checkpoint
            clear_checkpoint(st.session_state.project_name, "ta")
            st.session_state.pipeline_step = "ta_gen"
            st.rerun()

# ──── TA QA ────
elif step == "ta_qa":
    st.subheader("⚙️ Adım 2/3 — TA QA Hakem")
    if "ta_qa_result" not in st.session_state:
        from pipeline.brd.orchestrator import evaluate_ta_qa, finalize_stage
        from pipeline.brd.checkpoint import clear_checkpoint
        qa_model = st.session_state.get("qa_model")
        with st.status("🔍 TA QA değerlendiriyor...", expanded=True):
            qa = evaluate_ta_qa(st.session_state.ta_content, anthropic_key, gemini_key, log, model=qa_model)
        finalize_stage(st.session_state.run_id, "ta", st.session_state.ta_content, qa, st.session_state.get("ta_revision_count",0), False, 0)
        clear_checkpoint(st.session_state.project_name, "ta")
        st.session_state.ta_qa_result = qa
        st.session_state.ta_score = qa.get("genel_puan",0)

    qa = st.session_state.ta_qa_result
    score = st.session_state.ta_score
    show_qa_result(qa, "TA")
    show_log()

    st.divider()
    st.download_button("📥 TA Dokümanı İndir (.docx)",
        data=build_ta_docx(st.session_state.ta_content, st.session_state.project_name),
        file_name=f"{st.session_state.project_name}_teknik_analiz.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True)

    st.divider()
    if score >= TA_PASS_THRESHOLD:
        if st.button("➡️ Test Case'e Geç", type="primary", use_container_width=True):
            st.session_state.pipeline_step = "tc_gen"
            st.rerun()
    else:
        rev = st.session_state.get("ta_revision_count",0)
        if rev < MAX_REVISIONS:
            c1, c2 = st.columns(2)
            with c1:
                if st.button("🔄 Revize Et", type="primary", use_container_width=True, key="ta_rev"):
                    st.session_state.ta_revision_count = rev + 1
                    st.session_state.ta_feedback = qa.get("genel_degerlendirme","") + " | " + ", ".join(qa.get("iyilestirme_onerileri",[]))
                    del st.session_state["ta_qa_result"]
                    from pipeline.brd.checkpoint import clear_checkpoint
                    clear_checkpoint(st.session_state.project_name, "ta")
                    st.session_state.pipeline_step = "ta_gen"
                    st.rerun()
            with c2:
                if st.button("➡️ Yine de Devam Et", use_container_width=True, key="ta_fp"):
                    st.session_state.pipeline_step = "tc_gen"
                    st.rerun()
        else:
            st.warning("Max revizyon.")
            if st.button("➡️ Test Case'e Geç", type="primary", use_container_width=True):
                st.session_state.pipeline_step = "tc_gen"
                st.rerun()

# ──── TC GENERATE ────
elif step == "tc_gen":
    st.subheader("🧪 Adım 3/3 — Test Case Üretimi")
    from pipeline.brd.orchestrator import generate_tc
    gen_model = st.session_state.get("generation_model")
    with st.status("🤖 TC üretiliyor...", expanded=True) as s:
        tc = generate_tc(st.session_state.ba_content, st.session_state.ta_content, st.session_state.project_name,
                        st.session_state.jira_key, anthropic_key, gemini_key, log, st.session_state.get("tc_feedback",""), model=gen_model)
        s.update(label="✅ TC hazır!", state="complete")
    st.session_state.tc_content = tc
    st.session_state.pipeline_step = "tc_review"
    st.rerun()

# ──── TC REVIEW ────
elif step == "tc_review":
    st.subheader("🧪 Adım 3/3 — Test Case İnceleme")
    show_tc_preview(st.session_state.tc_content)
    show_log()
    st.divider()

    # Hakeme gönderme seçeneği
    skip_qa = st.checkbox("⚡ Hakeme göndermeden devam et (QA'yı atla)", key="skip_tc_qa")

    c1, c2 = st.columns(2)
    with c1:
        if st.button("✅ Onayla ve İlerle", type="primary", use_container_width=True, key="tc_ok"):
            if skip_qa:
                # QA'yı atla, direkt done'a geç - Database'e kaydet ve tamamla
                from pipeline.brd.orchestrator import finalize_stage, complete_run
                qa_result = {"genel_puan": 100, "gecti_mi": True, "genel_degerlendirme": "QA atlandı (Force Pass)",
                           "skorlar": [], "iyilestirme_onerileri": []}
                finalize_stage(st.session_state.run_id, "tc", st.session_state.tc_content, qa_result,
                             st.session_state.get("tc_revision_count", 0), True, 0)
                total_time = int(time.time() - st.session_state.get("pipeline_start", time.time()))
                complete_run(st.session_state.run_id, total_time)
                st.session_state.tc_score = 100  # Force pass
                st.session_state.tc_qa_result = qa_result
                st.session_state.pipeline_step = "done"
                st.success("⚡ QA atlandı, Pipeline tamamlandı!")
            else:
                # Normal akış, QA'ya gönder
                st.session_state.pipeline_step = "tc_qa"
            st.rerun()
    with c2:
        if st.button("🔄 Yeniden Üret", use_container_width=True, key="tc_re"):
            from pipeline.brd.checkpoint import clear_checkpoint
            clear_checkpoint(st.session_state.project_name, "tc")
            st.session_state.pipeline_step = "tc_gen"
            st.rerun()

# ──── TC QA ────
elif step == "tc_qa":
    st.subheader("🧪 Adım 3/3 — TC QA Hakem")
    if "tc_qa_result" not in st.session_state:
        from pipeline.brd.orchestrator import evaluate_tc_qa, finalize_stage, complete_run
        from pipeline.brd.checkpoint import clear_checkpoint
        qa_model = st.session_state.get("qa_model")
        with st.status("🔍 TC QA değerlendiriyor...", expanded=True):
            qa = evaluate_tc_qa(st.session_state.tc_content, anthropic_key, gemini_key, log, model=qa_model)
        finalize_stage(st.session_state.run_id, "tc", st.session_state.tc_content, qa, st.session_state.get("tc_revision_count",0), False, 0)
        clear_checkpoint(st.session_state.project_name, "tc")
        total_time = int(time.time() - st.session_state.get("pipeline_start", time.time()))
        complete_run(st.session_state.run_id, total_time)
        st.session_state.tc_qa_result = qa
        st.session_state.tc_score = qa.get("genel_puan",0)

    qa = st.session_state.tc_qa_result
    score = st.session_state.tc_score
    show_qa_result(qa, "TC")
    show_log()

    st.divider()
    # TC için Excel + CSV download
    tc_data = st.session_state.tc_content
    tcs = tc_data.get("test_cases", [])
    if tcs:
        import pandas as pd
        cols_23 = [
            "existance","created_by","date","app_bundle","test_case_id",
            "br_id","tr_id","priority","channel","testcase_type",
            "user_type","test_area","test_scenario","testcase","test_steps",
            "precondition","test_data","expected_result","postcondition",
            "actual_result","status","regression_case","comments"
        ]
        df = pd.DataFrame(tcs)
        for c in cols_23:
            if c not in df.columns:
                df[c] = ""
        df = df[cols_23]
        # Excel
        xlsx_buf = io.BytesIO()
        df.to_excel(xlsx_buf, index=False, sheet_name="Test Cases")
        xlsx_buf.seek(0)
        c1, c2 = st.columns(2)
        with c1:
            st.download_button("📥 TC İndir (.xlsx)",
                data=xlsx_buf,
                file_name=f"{st.session_state.project_name}_test_cases.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True)
        with c2:
            csv_data = df.to_csv(index=False, encoding="utf-8-sig")
            st.download_button("📥 TC İndir (.csv)",
                data=csv_data,
                file_name=f"{st.session_state.project_name}_test_cases.csv",
                mime="text/csv",
                use_container_width=True)

    st.divider()
    if score >= TC_PASS_THRESHOLD:
        if st.button("📊 Sonuçlara Git", type="primary", use_container_width=True):
            st.session_state.pipeline_step = "done"
            st.rerun()
    else:
        rev = st.session_state.get("tc_revision_count",0)
        if rev < MAX_REVISIONS:
            c1, c2 = st.columns(2)
            with c1:
                if st.button("🔄 Revize Et", type="primary", use_container_width=True, key="tc_rev"):
                    st.session_state.tc_revision_count = rev + 1
                    st.session_state.tc_feedback = qa.get("genel_degerlendirme","") + " | " + ", ".join(qa.get("iyilestirme_onerileri",[]))
                    del st.session_state["tc_qa_result"]
                    from pipeline.brd.checkpoint import clear_checkpoint
                    clear_checkpoint(st.session_state.project_name, "tc")
                    st.session_state.pipeline_step = "tc_gen"
                    st.rerun()
            with c2:
                if st.button("📊 Sonuçlara Git", use_container_width=True, key="tc_fp"):
                    st.session_state.pipeline_step = "done"
                    st.rerun()
        else:
            st.warning("Max revizyon.")
            if st.button("📊 Sonuçlara Git", type="primary", use_container_width=True):
                st.session_state.pipeline_step = "done"
                st.rerun()

# ──── DONE ────
elif step == "done":
    st.subheader("✅ Pipeline Tamamlandı!")
    show_log()
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("BA", f"{st.session_state.get('ba_score',0):.0f}/100")
    c2.metric("TA", f"{st.session_state.get('ta_score',0):.0f}/100")
    c3.metric("TC", f"{st.session_state.get('tc_score',0):.0f}/100")
    tc_n = len(st.session_state.get("tc_content",{}).get("test_cases",[]))
    c4.metric("TC Adet", tc_n)
    st.info("Detaylı sonuçlar ve export için **Pipeline Sonuç** sayfasına gidin.")
    if st.button("🔄 Yeni Pipeline Başlat", use_container_width=True):
        for k in list(st.session_state.keys()):
            if k not in ("anthropic_api_key","gemini_api_key"):
                del st.session_state[k]
        st.rerun()
