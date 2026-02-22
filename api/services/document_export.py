"""
Document Export Service - Convert BA/TA/TC JSON to DOCX/XLSX formats
Uses the same format as Streamlit's build functions.
"""
from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from io import BytesIO
from typing import Dict, Any
import json
import time
import pandas as pd


def json_to_ba_docx(ba_content: Dict[str, Any]) -> BytesIO:
    """
    Convert BA JSON to DOCX format (matching Streamlit's build_ba_docx).

    Expected structure:
    {
        "ekranlar": [
            {
                "ekran_adi": str,
                "aciklama": str,
                "gerekli_dokumanlar": {"teknik_akis": str, "tasarim_dosyasi": str},
                "is_akisi_diyagrami": [str],
                "fonksiyonel_gereksinimler": [{"id": str, "tanim": str}],
                "is_kurallari": [{"kural": str, "detay": str}],
                "kabul_kriterleri": [{"id": str, "kriter": str}],
                "validasyonlar": [{"alan": str, "kisit": str, "hata_mesaji": str}]
            }
        ]
    }
    """
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)

    # Title
    project_name = ba_content.get("_project_name", "Project")
    doc.add_heading(f'{project_name} — İş Analizi', level=0)
    doc.add_paragraph(f'Tarih: {time.strftime("%d.%m.%Y %H:%M")}')

    # Process each screen
    for i, e in enumerate(ba_content.get("ekranlar", []), 1):
        # 1. Geliştirmenin Adı
        doc.add_heading(f'{i}. {e.get("ekran_adi", "")}', level=1)

        # 1.1. Açıklama
        doc.add_heading(f'{i}.1. Açıklama', level=2)
        doc.add_paragraph(e.get("aciklama", ""))

        # 1.2. Gerekli Dokümanlar
        doc.add_heading(f'{i}.2. Gerekli Dokümanlar', level=2)
        gerekli_dok = e.get("gerekli_dokumanlar", {})
        doc.add_heading(f'{i}.2.1. Teknik Akış/Analiz', level=3)
        doc.add_paragraph(gerekli_dok.get("teknik_akis", "") if isinstance(gerekli_dok, dict) else "")
        doc.add_heading(f'{i}.2.2. Tasarım Dosyası', level=3)
        doc.add_paragraph(gerekli_dok.get("tasarim_dosyasi", "") if isinstance(gerekli_dok, dict) else "")

        # 1.3. İş Akışı Diyagramı
        doc.add_heading(f'{i}.3. İş Akışı Diyagramı', level=2)
        for adim in e.get("is_akisi_diyagrami", []):
            doc.add_paragraph(adim, style='List Number')

        # 1.4. Fonksiyonel Gereksinimler
        frs = e.get("fonksiyonel_gereksinimler", [])
        doc.add_heading(f'{i}.4. Fonksiyonel Gereksinimler', level=2)
        for fr in frs:
            p = doc.add_paragraph()
            run = p.add_run(f'{fr.get("id","")}: ')
            run.bold = True
            p.add_run(fr.get("tanim", ""))

        # 1.5. İş Kuralları
        kurallar = e.get("is_kurallari", [])
        doc.add_heading(f'{i}.5. İş Kuralları', level=2)
        for k in kurallar:
            p = doc.add_paragraph()
            run = p.add_run(k.get("kural", ""))
            run.bold = True
            if k.get("detay"):
                doc.add_paragraph(k["detay"])

        # 1.6. Kabul Kriterleri
        brs = e.get("kabul_kriterleri", [])
        doc.add_heading(f'{i}.6. Kabul Kriterleri', level=2)
        for br in brs:
            p = doc.add_paragraph()
            run = p.add_run(f'{br.get("id","")}: ')
            run.bold = True
            p.add_run(br.get("kriter", ""))

        # 1.7. Validasyonlar
        vals = e.get("validasyonlar", [])
        doc.add_heading(f'{i}.7. Validasyonlar', level=2)
        for v in vals:
            doc.add_paragraph(
                f'{v.get("alan","")}: {v.get("kisit","")} → {v.get("hata_mesaji","")}',
                style='List Bullet'
            )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def json_to_ta_docx(ta_content: Dict[str, Any]) -> BytesIO:
    """
    Convert TA JSON to DOCX format (matching Streamlit's build_ta_docx).

    Expected structure:
    {
        "teknik_analiz": {
            "genel_tanim": {"modul_adi": str, "teknoloji_stack": [str], "mimari_yaklasim": str},
            "endpoint_detaylari": {
                "/api/path": {
                    "method": str,
                    "aciklama": str,
                    "request_body": dict,
                    "response_success": dict,
                    "response_errors": [{"http_code": str, "error_code": str, "mesaj": str}]
                }
            },
            "dto_veri_yapilari": [
                {
                    "dto_adi": str,
                    "aciklama": str,
                    "fields": [{"field": str, "tip": str, "validasyon": str}]
                }
            ],
            "validasyon_kurallari": [{"id": str, "field": str, "kural": str}],
            "mock_curl_ornekleri": [{"endpoint_adi": str, "curl": str}]
        }
    }
    """
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)

    ta = ta_content.get("teknik_analiz", ta_content)

    # Title
    project_name = ta_content.get("_project_name", "Project")
    doc.add_heading(f'{project_name} — Teknik Analiz', level=0)
    doc.add_paragraph(f'Tarih: {time.strftime("%d.%m.%Y %H:%M")}')

    # Genel Tanım — handle both mock and AI formats
    genel = ta.get("genel_tanim", {})
    if genel:
        doc.add_heading("Genel Tanım", level=1)
        doc.add_paragraph(f'Modül: {genel.get("modul_adi","")}')
        stack = genel.get("teknoloji_stack", [])
        if isinstance(stack, list):
            doc.add_paragraph(f'Stack: {", ".join(stack)}')
        else:
            doc.add_paragraph(f'Stack: {stack}')
        doc.add_paragraph(f'Mimari: {genel.get("mimari_yaklasim","")}')

    # API Endpoints — handle both dict format (mock) and list format (AI output)
    endpoints_dict = ta.get("endpoint_detaylari", {})
    endpoints_list = ta.get("api_endpoints", [])

    if endpoints_dict and isinstance(endpoints_dict, dict):
        doc.add_heading("API Endpoints", level=1)
        for ep, d in endpoints_dict.items():
            doc.add_heading(f'{d.get("method","GET")} {ep}', level=2)
            doc.add_paragraph(d.get("aciklama", ""))
            if d.get("request_body"):
                doc.add_paragraph(f'Request: {json.dumps(d["request_body"], ensure_ascii=False, indent=2)[:500]}')
            if d.get("response_success"):
                doc.add_paragraph(f'Response: {json.dumps(d["response_success"], ensure_ascii=False, indent=2)[:500]}')
            for err in d.get("response_errors", []):
                doc.add_paragraph(f'{err.get("http_code","")} {err.get("error_code","")}: {err.get("mesaj","")}', style='List Bullet')
    elif endpoints_list and isinstance(endpoints_list, list):
        doc.add_heading("API Endpoints", level=1)
        for ep in endpoints_list:
            method = ep.get("method", ep.get("http_method", "GET"))
            path = ep.get("endpoint", ep.get("path", ep.get("url", "")))
            doc.add_heading(f'{method} {path}', level=2)
            doc.add_paragraph(ep.get("aciklama", ep.get("description", "")))
            if ep.get("request_body"):
                body = ep["request_body"]
                doc.add_paragraph(f'Request: {json.dumps(body, ensure_ascii=False, indent=2)[:500]}')
            if ep.get("response_success", ep.get("response", None)):
                resp = ep.get("response_success", ep.get("response", {}))
                doc.add_paragraph(f'Response: {json.dumps(resp, ensure_ascii=False, indent=2)[:500]}')
            for err in ep.get("response_errors", ep.get("error_responses", [])):
                doc.add_paragraph(f'{err.get("http_code", err.get("status_code", ""))} {err.get("error_code","")}: {err.get("mesaj", err.get("message", ""))}', style='List Bullet')

    # DTO Yapıları — handle both "dto_veri_yapilari" (mock) and "dtos" (AI)
    dtos = ta.get("dto_veri_yapilari", ta.get("dtos", []))
    if dtos:
        doc.add_heading("DTO Yapıları", level=1)
        for dto in dtos:
            dto_name = dto.get("dto_adi", dto.get("dto_name", dto.get("name", "")))
            doc.add_heading(dto_name, level=2)
            doc.add_paragraph(dto.get("aciklama", dto.get("description", "")))
            for f in dto.get("fields", []):
                field_name = f.get("field", f.get("field_name", f.get("name", "")))
                field_type = f.get("tip", f.get("data_type", f.get("type", "")))
                field_val = f.get("validasyon", f.get("validation", ""))
                doc.add_paragraph(f'{field_name}: {field_type} {field_val}', style='List Bullet')

    # Validasyon Kuralları
    vrs = ta.get("validasyon_kurallari", ta.get("validation_rules", []))
    if vrs:
        doc.add_heading("Validasyon Kuralları", level=1)
        for v in vrs:
            doc.add_paragraph(f'{v.get("id","")} [{v.get("field","")}]: {v.get("kural", v.get("rule", ""))}', style='List Bullet')

    # cURL Örnekleri
    curls = ta.get("mock_curl_ornekleri", ta.get("curl_ornekleri", []))
    if curls:
        doc.add_heading("cURL Örnekleri", level=1)
        for c in curls:
            doc.add_paragraph(c.get("endpoint_adi", c.get("endpoint_name", "")), style='List Bullet')
            doc.add_paragraph(c.get("curl", ""))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def json_to_tc_xlsx(tc_content: Dict[str, Any]) -> BytesIO:
    """
    Convert TC JSON to XLSX in 23-column Loodos template format.

    Matches the exact format used in Streamlit utils/export.py.

    Expected structure:
    {
        "test_cases": [
            {
                "test_case_id": str,
                "br_id": str,
                "priority": str,
                "test_area": str,
                "testcase": str,
                "test_steps": str,
                "test_data": str,
                "expected_result": str,
                // Plus optional fields: existance, created_by, date, app_bundle,
                // tr_id, channel, testcase_type, user_type, test_scenario,
                // precondition, postcondition, actual_result, status,
                // regression_case, comments
            }
        ]
    }
    """
    # Get test cases
    test_cases = tc_content.get("test_cases", [])

    # Create workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Test Cases"

    # 23-column Loodos headers
    headers = [
        "EXISTANCE", "CREATED BY", "DATE", "APP BUNDLE", "TEST CASE ID",
        "BR ID", "TR ID", "PRIORITY", "CHANNEL", "TESTCASE TYPE",
        "USER TYPE", "TEST AREA", "TEST SCENARIO", "TESTCASE", "TEST STEPS",
        "PRECONDITION", "TEST DATA", "EXPECTED RESULT", "POSTCONDITION",
        "ACTUAL RESULT", "STATUS", "REGRESSION CASE", "COMMENTS"
    ]

    # Styling (Loodos template colors)
    header_fill = PatternFill(start_color="1a1d2e", end_color="1a1d2e", fill_type="solid")
    header_font = Font(name="Calibri", bold=True, color="FFFFFF", size=10)
    thin_border = Border(
        left=Side(style="thin", color="555555"),
        right=Side(style="thin", color="555555"),
        top=Side(style="thin", color="555555"),
        bottom=Side(style="thin", color="555555"),
    )

    # Write headers
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = thin_border

    # Field mapping (JSON keys to column numbers)
    field_map = {
        "existance": 1, "created_by": 2, "date": 3, "app_bundle": 4,
        "test_case_id": 5, "br_id": 6, "tr_id": 7, "priority": 8,
        "channel": 9, "testcase_type": 10, "user_type": 11, "test_area": 12,
        "test_scenario": 13, "testcase": 14, "test_steps": 15,
        "precondition": 16, "test_data": 17, "expected_result": 18,
        "postcondition": 19, "actual_result": 20, "status": 21,
        "regression_case": 22, "comments": 23,
    }

    # Write data rows
    for row_idx, tc in enumerate(test_cases, 2):
        for field, col in field_map.items():
            val = tc.get(field, "")
            cell = ws.cell(row=row_idx, column=col, value=str(val))
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            cell.border = thin_border

    # Column widths (optimized for 23 columns)
    widths = [10, 12, 10, 15, 20, 10, 10, 10, 8, 12, 12, 15, 25, 30, 40, 20, 20, 30, 15, 15, 10, 12, 15]
    for i, width in enumerate(widths, 1):
        col_letter = chr(64 + i) if i <= 26 else "A" + chr(64 + i - 26)
        ws.column_dimensions[col_letter].width = width

    # Auto-filter and freeze panes
    ws.auto_filter.ref = f"A1:W{len(test_cases) + 1}"
    ws.freeze_panes = "A2"

    # Save to BytesIO
    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer
