"""
Document Reader — Parse DOCX files to structured JSON
"""
from docx import Document
from typing import Dict, Any, List
import json


def parse_docx_to_json(file_path: str) -> Dict[str, Any]:
    """
    Parse a DOCX file and convert to structured JSON.

    Attempts to extract structured data based on document type (BA/TA/TC).
    Falls back to paragraph-based structure if specific format not detected.

    Args:
        file_path: Path to DOCX file

    Returns:
        Structured JSON dictionary
    """
    try:
        doc = Document(file_path)

        # Extract all text content
        paragraphs = []
        tables = []

        # Get paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append({
                    'text': para.text.strip(),
                    'style': para.style.name if para.style else 'Normal'
                })

        # Get tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                tables.append(table_data)

        # Try to detect document type and structure accordingly
        full_text = '\n'.join([p['text'] for p in paragraphs])

        # Check if it's a TA document
        if any(keyword in full_text.lower() for keyword in ['teknik analiz', 'endpoint', 'api', 'dto', 'mimari']):
            return parse_ta_document(paragraphs, tables, full_text)

        # Check if it's a BA document
        elif any(keyword in full_text.lower() for keyword in ['i̇ş analizi', 'ekran', 'fonksiyonel gereksinim', 'kabul kriteri']):
            return parse_ba_document(paragraphs, tables, full_text)

        # Check if it's a TC document
        elif any(keyword in full_text.lower() for keyword in ['test case', 'test adımları', 'beklenen sonuç']):
            return parse_tc_document(paragraphs, tables, full_text)

        # Generic structure
        return {
            'paragraphs': paragraphs,
            'tables': tables,
            'full_text': full_text,
            'document_type': 'unknown'
        }

    except Exception as e:
        raise ValueError(f"DOCX parsing failed: {str(e)}")


def parse_ta_document(paragraphs: List[Dict], tables: List[List], full_text: str) -> Dict[str, Any]:
    """Parse Technical Analysis document."""

    # Simple TA structure extraction
    ta_structure = {
        'teknik_analiz': {
            'genel_tanim': {},
            'endpoint_detaylari': {},
            'dto_veri_yapilari': [],
            'validasyon_kurallari': [],
            'mock_curl_ornekleri': []
        },
        'document_type': 'ta',
        'raw_paragraphs': paragraphs,
        'raw_tables': tables
    }

    # Try to extract module name from first heading
    for para in paragraphs[:5]:
        if para['style'] in ['Heading 1', 'Title'] and para['text']:
            ta_structure['teknik_analiz']['genel_tanim']['modul_adi'] = para['text']
            break

    # Extract endpoints from tables or text
    current_section = None
    for para in paragraphs:
        text_lower = para['text'].lower()

        # Detect sections
        if 'endpoint' in text_lower or 'api' in text_lower:
            current_section = 'endpoints'
        elif 'dto' in text_lower or 'veri yapı' in text_lower:
            current_section = 'dto'
        elif 'validasyon' in text_lower:
            current_section = 'validation'

        # Extract content based on section
        if current_section == 'endpoints' and para['text'].startswith('/'):
            # Looks like an endpoint path
            endpoint_path = para['text'].split()[0]
            ta_structure['teknik_analiz']['endpoint_detaylari'][endpoint_path] = {
                'method': 'GET',
                'aciklama': para['text']
            }

    return ta_structure


def parse_ba_document(paragraphs: List[Dict], tables: List[List], full_text: str) -> Dict[str, Any]:
    """Parse Business Analysis document."""

    ba_structure = {
        'ekranlar': [],
        'document_type': 'ba',
        'raw_paragraphs': paragraphs,
        'raw_tables': tables
    }

    # Try to extract screens
    current_screen = None
    for para in paragraphs:
        if para['style'] in ['Heading 1', 'Heading 2']:
            # New screen
            if 'ekran' in para['text'].lower() or 'sayfa' in para['text'].lower():
                if current_screen:
                    ba_structure['ekranlar'].append(current_screen)
                current_screen = {
                    'ekran_adi': para['text'],
                    'aciklama': '',
                    'fonksiyonel_gereksinimler': [],
                    'is_kurallari': [],
                    'kabul_kriterleri': [],
                    'validasyonlar': []
                }
        elif current_screen and para['text']:
            # Add to description
            if not current_screen['aciklama']:
                current_screen['aciklama'] = para['text']

    if current_screen:
        ba_structure['ekranlar'].append(current_screen)

    return ba_structure


def parse_tc_document(paragraphs: List[Dict], tables: List[List], full_text: str) -> Dict[str, Any]:
    """Parse Test Case document."""

    tc_structure = {
        'test_cases': [],
        'document_type': 'tc',
        'raw_paragraphs': paragraphs,
        'raw_tables': tables
    }

    # Extract test cases from tables
    for table_idx, table in enumerate(tables):
        if len(table) > 1:  # Has header + data
            # Try to identify TC table structure
            header = [cell.lower() for cell in table[0]]

            # Check if it's a test case table
            if any('test' in h or 'tc' in h for h in header):
                for row in table[1:]:
                    if len(row) >= 3:  # Minimum: ID, case, expected
                        tc_structure['test_cases'].append({
                            'test_case_id': row[0] if len(row) > 0 else '',
                            'testcase': row[1] if len(row) > 1 else '',
                            'expected_result': row[2] if len(row) > 2 else '',
                            'test_steps': row[3] if len(row) > 3 else '',
                            'priority': row[4] if len(row) > 4 else 'Medium',
                            'test_area': row[5] if len(row) > 5 else '',
                        })

    return tc_structure
