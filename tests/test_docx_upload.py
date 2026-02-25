"""
Test DOCX upload functionality
"""
from playwright.sync_api import sync_playwright
import time
from docx import Document

# First, create a sample DOCX file
print("Creating sample BRD.docx file...")
doc = Document()
doc.add_heading('Kullanıcı Giriş Sistemi', 0)

doc.add_heading('1. Amaç', level=1)
doc.add_paragraph('Bu doküman, web uygulaması için kullanıcı giriş (login) sisteminin iş analizini içermektedir.')

doc.add_heading('2. Kapsam', level=1)
doc.add_paragraph('- Kullanıcı giriş ekranı')
doc.add_paragraph('- Şifre doğrulama')
doc.add_paragraph('- "Beni Hatırla" özelliği')
doc.add_paragraph('- Şifremi Unuttum akışı')

doc.add_heading('3. Fonksiyonel Analiz', level=1)
doc.add_heading('3.1 Login Ekranı', level=2)
doc.add_paragraph('Ekran Adı: Login')
doc.add_paragraph('Açıklama: Kullanıcının email ve şifre ile sisteme giriş yapabileceği ekran')

doc.add_heading('Form Alanları:', level=3)
doc.add_paragraph('- Email (zorunlu, email formatı)')
doc.add_paragraph('- Şifre (zorunlu, min 8 karakter)')
doc.add_paragraph('- Beni Hatırla (checkbox, opsiyonel)')

doc.add_heading('İş Kuralları:', level=3)
doc.add_paragraph('1. Email formatı kontrol edilmeli (@domain.com)')
doc.add_paragraph('2. Şifre minimum 8 karakter olmalı')
doc.add_paragraph('3. 3 başarısız denemeden sonra hesap 15 dakika kilitlenmeli')

doc.save('/tmp/sample_brd.docx')
print("✓ Created: /tmp/sample_brd.docx")

# Now test uploading it
print("\nTesting DOCX upload...")

with sync_playwright() as p:
    browser = p.chromium.launch(headless=False, slow_mo=300)
    page = browser.new_page()

    print("1. Opening BRD Pipeline page...")
    page.goto('http://localhost:5173/brd-pipeline')
    page.wait_for_load_state('networkidle')
    time.sleep(1)

    print("2. Selecting project...")
    page.locator('.ant-select-selector').first.click()
    time.sleep(0.5)
    page.locator('.ant-select-item').first.click()
    print("   ✓ Project selected")

    print("3. Uploading DOCX file...")

    # Click on upload button to trigger file chooser
    page.locator('input[type="file"]').set_input_files('/tmp/sample_brd.docx')

    # Wait for file to be processed
    print("   ⏳ Waiting for DOCX to be parsed...")
    time.sleep(3)

    # Check if content was filled
    textarea_content = page.locator('textarea[placeholder*="BRD"]').input_value()

    if textarea_content and len(textarea_content) > 100:
        print(f"   ✓ DOCX parsed successfully! ({len(textarea_content)} characters)")
        print(f"\n   Preview (first 200 chars):")
        print(f"   {textarea_content[:200]}...")
    else:
        print("   ✗ DOCX parsing failed or content is empty")
        print(f"   Content length: {len(textarea_content)}")

    # Take screenshot
    page.screenshot(path='/tmp/docx_upload_test.png', full_page=True)
    print("\n4. Screenshot saved: /tmp/docx_upload_test.png")

    print("\n5. Keeping browser open for 5 seconds...")
    time.sleep(5)

    browser.close()

print("\n✅ Test completed!")
print("Check the screenshot at: /tmp/docx_upload_test.png")
