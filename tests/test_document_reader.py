"""
Test Document Reader - Word and Google Drive support

This test demonstrates the new document import features:
1. Reading Word (.docx) documents
2. Downloading from Google Drive

Run this after installing dependencies:
pip install python-docx requests
"""

def test_google_drive_url_parsing():
    """Test Google Drive URL parsing"""
    from pipeline.document_reader import extract_google_drive_file_id

    print("🧪 Testing Google Drive URL Parsing\n")
    print("=" * 70)

    test_cases = [
        ("https://drive.google.com/file/d/1ABC123xyz/view", "1ABC123xyz"),
        ("https://docs.google.com/document/d/1XYZ789abc/edit", "1XYZ789abc"),
        ("https://drive.google.com/open?id=1DEF456pqr", "1DEF456pqr"),
        ("https://docs.google.com/spreadsheets/d/1GHI789stu/edit#gid=0", "1GHI789stu"),
    ]

    all_passed = True
    for url, expected_id in test_cases:
        file_id = extract_google_drive_file_id(url)
        status = "✅" if file_id == expected_id else "❌"
        print(f"{status} URL: {url}")
        print(f"   Expected: {expected_id}")
        print(f"   Got:      {file_id}")
        print()

        if file_id != expected_id:
            all_passed = False

    if all_passed:
        print("🎉 All URL parsing tests passed!")
    else:
        print("⚠️  Some tests failed")

    return all_passed


def test_word_document_reading():
    """Test Word document reading"""
    from docx import Document
    from pipeline.document_reader import read_docx
    import io

    print("\n🧪 Testing Word Document Reading\n")
    print("=" * 70)

    # Create a test document in memory
    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_heading('Section 1', 1)
    doc.add_paragraph('This is a paragraph in section 1.')
    doc.add_paragraph('• Bullet point 1', style='List Bullet')
    doc.add_paragraph('• Bullet point 2', style='List Bullet')

    doc.add_heading('Subsection 1.1', 2)
    doc.add_paragraph('Content in subsection.')

    doc.add_heading('Section 2', 1)
    doc.add_paragraph('Content in section 2.')

    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)

    # Read with our reader
    extracted_text = read_docx(doc_bytes.read())

    print("Extracted Text:")
    print("-" * 70)
    print(extracted_text)
    print("-" * 70)

    # Check if headings are marked
    checks = [
        ("# Test Document" in extracted_text, "Title (H1) marked with #"),
        ("# Section 1" in extracted_text, "Heading 1 marked with #"),
        ("## Subsection 1.1" in extracted_text, "Heading 2 marked with ##"),
        ("Bullet point" in extracted_text, "Bullet points preserved"),
    ]

    print("\nChecks:")
    all_passed = True
    for passed, description in checks:
        status = "✅" if passed else "❌"
        print(f"{status} {description}")
        if not passed:
            all_passed = False

    if all_passed:
        print("\n🎉 Word document reading works correctly!")
    else:
        print("\n⚠️  Some checks failed")

    return all_passed


def test_integration():
    """Test full integration with parser"""
    from pipeline.document_reader import read_docx
    from pipeline.document_parser_v2 import parse_text_to_json
    from docx import Document
    import io
    import json

    print("\n🧪 Testing Full Integration (Word → Parser → JSON)\n")
    print("=" * 70)

    # Create a test BA document
    doc = Document()
    doc.add_heading('User Login - Business Analysis', 0)

    doc.add_heading('Kullanıcı Arayüzleri', 1)
    doc.add_heading('Giriş Ekranı', 2)
    doc.add_paragraph('• Email: Kullanıcı email adresi (zorunlu)', style='List Bullet')
    doc.add_paragraph('• Şifre: Kullanıcı şifresi (zorunlu)', style='List Bullet')

    doc.add_heading('Backend İşlemleri', 1)
    doc.add_heading('User Login', 2)
    doc.add_paragraph('POST /api/auth/login')
    doc.add_paragraph('Kullanıcı girişi yapar')

    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)

    # Step 1: Extract text
    print("Step 1: Extracting text from Word document...")
    extracted_text = read_docx(doc_bytes.read())
    print(f"✅ Extracted {len(extracted_text)} characters")

    # Step 2: Parse with rule-based parser
    print("\nStep 2: Parsing with rule-based parser...")
    parsed_json = parse_text_to_json(extracted_text, 'ba')

    print(f"✅ Parsed JSON:")
    print(json.dumps(parsed_json, indent=2, ensure_ascii=False))

    # Verify results
    print("\nVerification:")
    checks = [
        (len(parsed_json.get('ekranlar', [])) > 0, "Screens (ekranlar) found"),
        (len(parsed_json.get('backend_islemler', [])) > 0, "Backend operations found"),
    ]

    all_passed = True
    for passed, description in checks:
        status = "✅" if passed else "❌"
        print(f"{status} {description}")
        if not passed:
            all_passed = False

    if all_passed:
        print("\n🎉 Full integration works! Word → Parser → JSON ✅")
    else:
        print("\n⚠️  Integration test failed")

    return all_passed


if __name__ == "__main__":
    print("📄 Document Reader Test Suite")
    print("=" * 70)
    print()

    try:
        # Test 1: URL parsing (no external dependencies)
        result1 = test_google_drive_url_parsing()

        # Test 2: Word document reading
        try:
            result2 = test_word_document_reading()
        except ImportError:
            print("\n⚠️  python-docx not installed, skipping Word tests")
            print("   Install with: pip install python-docx")
            result2 = None

        # Test 3: Full integration
        if result2 is not None:
            try:
                result3 = test_integration()
            except Exception as e:
                print(f"\n❌ Integration test failed: {e}")
                result3 = False
        else:
            result3 = None

        # Summary
        print("\n" + "=" * 70)
        print("📊 Test Summary")
        print("=" * 70)
        print(f"Google Drive URL Parsing: {'✅ PASS' if result1 else '❌ FAIL'}")
        print(f"Word Document Reading:    {'✅ PASS' if result2 else '⚠️  SKIPPED' if result2 is None else '❌ FAIL'}")
        print(f"Full Integration:         {'✅ PASS' if result3 else '⚠️  SKIPPED' if result3 is None else '❌ FAIL'}")

    except ImportError as e:
        print(f"\n❌ Import Error: {e}")
        print("\nRequired dependencies:")
        print("  pip install python-docx requests")
    except Exception as e:
        print(f"\n❌ Unexpected error: {e}")
        import traceback
        traceback.print_exc()
