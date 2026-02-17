"""
Pytest configuration and shared fixtures for BA-QA Platform tests.
"""
import io
import pytest
import sqlite3
import tempfile
import os
from pathlib import Path


# ─────────────────────────────────────────────
# Database Fixtures
# ─────────────────────────────────────────────

@pytest.fixture
def temp_db():
    """Create a temporary test database."""
    # Create a temporary file for the database
    db_fd, db_path = tempfile.mkstemp(suffix='.db')

    yield db_path

    # Cleanup
    os.close(db_fd)
    os.unlink(db_path)


@pytest.fixture
def db_connection(temp_db):
    """Create a database connection with schema initialized."""
    # Import here to avoid circular imports
    from data.database import init_db, get_db

    # Temporarily override DB_PATH
    import data.database as db_module
    original_db_path = db_module.DB_PATH
    db_module.DB_PATH = temp_db

    # Initialize database
    init_db()

    # Get connection
    conn = get_db()

    yield conn

    # Cleanup
    conn.close()
    db_module.DB_PATH = original_db_path


@pytest.fixture
def clean_db(db_connection):
    """Provide a clean database for each test."""
    yield db_connection

    # Rollback any changes after test
    db_connection.rollback()


# ─────────────────────────────────────────────
# Sample Data Fixtures
# ─────────────────────────────────────────────

@pytest.fixture
def sample_project():
    """Sample project data."""
    return {
        "name": "Test Mobile App",
        "description": "Test project for mobile application",
        "jira_project_key": "TMA",
        "tags": ["mobile", "test", "ios"]
    }


@pytest.fixture
def sample_ba_document():
    """Sample BA document content."""
    return {
        "ekranlar": [
            {
                "ekran_adi": "Login Screen",
                "aciklama": "User authentication",
                "ui_elementleri": ["Email input", "Password input", "Login button"]
            },
            {
                "ekran_adi": "Dashboard",
                "aciklama": "Main dashboard",
                "ui_elementleri": ["Account balance", "Recent transactions"]
            }
        ],
        "backend_islemler": [
            {
                "islem_adi": "User Authentication",
                "endpoint": "/api/auth/login",
                "method": "POST"
            }
        ]
    }


@pytest.fixture
def sample_tc_document():
    """Sample TC document content."""
    return {
        "test_scenarios": [
            {
                "scenario_id": "TC001",
                "title": "Successful Login",
                "steps": [
                    "Open login page",
                    "Enter valid credentials",
                    "Click login button"
                ],
                "expected_result": "User is redirected to dashboard"
            }
        ]
    }


# ─────────────────────────────────────────────
# API Mocking Fixtures
# ─────────────────────────────────────────────

@pytest.fixture
def mock_anthropic_response():
    """Mock Anthropic API response."""
    return {
        "json_output": {
            "ekranlar": [{"ekran_adi": "Test Screen"}],
            "backend_islemler": []
        },
        "stop_reason": "end_turn"
    }


@pytest.fixture
def mock_gemini_response():
    """Mock Gemini API response."""
    return {
        "json_output": {
            "genel_puan": 85.0,
            "gecti_mi": True,
            "eksikler": []
        },
        "stop_reason": "stop"
    }


# ─────────────────────────────────────────────
# Test Markers
# ─────────────────────────────────────────────

def pytest_configure(config):
    """Configure custom markers."""
    config.addinivalue_line(
        "markers", "unit: Unit tests (fast, isolated)"
    )
    config.addinivalue_line(
        "markers", "integration: Integration tests (slower, with dependencies)"
    )
    config.addinivalue_line(
        "markers", "slow: Slow running tests"
    )
    config.addinivalue_line(
        "markers", "db: Tests that require database"
    )
    config.addinivalue_line(
        "markers", "api: Tests that require API calls"
    )


# ─────────────────────────────────────────────
# DOCX Fixtures — Adım 5
# ─────────────────────────────────────────────

def _add_hyperlink_to_para(para, url: str, text: str):
    """python-docx ile paragraf içine hyperlink ekle."""
    from docx.oxml import OxmlElement
    part  = para.part
    r_id  = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(
        '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id', r_id
    )
    run_el = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    run_el.append(t)
    hyperlink.append(run_el)
    para._element.append(hyperlink)


def _add_bullet_item(doc, text: str, level: int = 0) -> None:
    """Belirtilen level'da bullet list item ekle (numPr XML ile)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p = doc.add_paragraph(text)
    pPr = OxmlElement('w:pPr')
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(level))
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), '1')
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)
    p._element.insert(0, pPr)


@pytest.fixture(scope="session")
def sample_loodos_ba_docx() -> bytes:
    """
    Kahve Dünyası stilinde gerçekçi bir Loodos BA DOCX dosyası oluşturur.

    Yapı:
      H1: Proje Açıklaması   → platform/dil bilgisi
      H1: Proje Kapsamı      → kapsam metni
      H1: Mobil Uygulama Gereksinimleri
        H2: Splash
          H3: Açıklama        → açıklama paragrafı
          H3: Tasarım Dosyaları → Figma linki
          H3: İş Akışı        → L0/L1/L2 nested bullet + Lottie link
        H2: Login
          H3: Açıklama
          H3: İş Akışı        → L0/L1 bullet, bold run
          H3: OTP Akışı       → ayrı is_akisi

    Returns:
        bytes — .docx binary içeriği
    """
    from docx import Document

    doc = Document()

    # ── Proje Açıklaması ──────────────────────────────────────
    doc.add_heading('Proje Açıklaması', level=1)
    doc.add_paragraph('Bu projede amaç; Kahve Dünyası mobil uygulamasını geliştirmektir.')
    doc.add_paragraph('Platform Bilgisi: iOS | Android | Web')
    doc.add_paragraph('Dil Desteği: Türkçe, İngilizce')

    # ── Proje Kapsamı ─────────────────────────────────────────
    doc.add_heading('Proje Kapsamı', level=1)
    doc.add_paragraph('Mobil ve Web tarafında yapılan geliştirmeler kapsam dahilindedir.')

    # ── Mobil Uygulama Gereksinimleri ─────────────────────────
    doc.add_heading('Mobil Uygulama Gereksinimleri', level=1)

    # ── Splash ekranı ────────────────────────────────────────
    doc.add_heading('Splash', level=2)

    doc.add_heading('Açıklama', level=3)
    doc.add_paragraph('Uygulama başlangıç ekranıdır.')

    doc.add_heading('Tasarım Dosyaları', level=3)
    figma_url = 'https://figma.com/file/splash-design/KahveDunyasi'
    para = doc.add_paragraph('Figma: ')
    _add_hyperlink_to_para(para, figma_url, 'Splash Tasarımı')

    doc.add_heading('İş Akışı', level=3)
    lottie_url = 'https://app.lottiefiles.com/share/splash-animation'
    _add_bullet_item(doc, 'Kullanıcı uygulamayı açar; splash ekranı görüntülenir.', level=0)
    _add_bullet_item(doc, 'Kahve Dünyası logosu ve animasyon gösterilir.', level=1)
    para_lottie = doc.add_paragraph()
    para_lottie_content = 'Lottie animasyon dosyası: '
    _add_hyperlink_to_para(para_lottie, lottie_url, 'Lottie Linki')
    _add_bullet_item(doc, 'Al götür logosu gösterilir.', level=1)
    _add_bullet_item(doc, '2 saniye beklenir, ardından Login ekranına geçilir.', level=0)

    # ── Login ekranı ─────────────────────────────────────────
    doc.add_heading('Login', level=2)

    doc.add_heading('Açıklama', level=3)
    doc.add_paragraph('Kullanıcı telefon numarası ile sisteme giriş yapar.')

    doc.add_heading('İş Akışı', level=3)
    _add_bullet_item(doc, 'Kullanıcı telefon numarasını girer.', level=0)
    # Bold run ekle
    p_bold = doc.add_paragraph()
    run_norm = p_bold.add_run('Giriş Yap ')
    run_bold = p_bold.add_run('butonuna')
    run_bold.bold = True
    run_norm2 = p_bold.add_run(' tıklar.')
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = OxmlElement('w:pPr')
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl'); ilvl.set(qn('w:val'), '1')
    numId = OxmlElement('w:numId'); numId.set(qn('w:val'), '1')
    numPr.append(ilvl); numPr.append(numId); pPr.append(numPr)
    p_bold._element.insert(0, pPr)
    _add_bullet_item(doc, 'OTP doğrulama ekranına yönlendirilir.', level=0)

    doc.add_heading('OTP Akışı', level=3)
    _add_bullet_item(doc, 'Kullanıcıya 6 haneli OTP kodu SMS ile gönderilir.', level=0)
    _add_bullet_item(doc, 'Kod 3 dakika geçerlidir.', level=1)
    _add_bullet_item(doc, 'Hatalı giriş 3 kez yapılırsa hesap kilitlenir.', level=1)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture(scope="session")
def sample_loodos_ba_elements(sample_loodos_ba_docx) -> list:
    """
    sample_loodos_ba_docx'ten read_docx_structured() çıktısı döndürür.
    BADocxParser ve Orchestrator testlerinde kullanılabilir.
    """
    from pipeline.document_reader import read_docx_structured
    return read_docx_structured(sample_loodos_ba_docx)


@pytest.fixture(scope="session")
def sample_loodos_ba_parsed(sample_loodos_ba_elements) -> dict:
    """
    sample_loodos_ba_elements'ten BADocxParser().parse() çıktısı döndürür.
    """
    from pipeline.ba_docx_parser import BADocxParser
    return BADocxParser().parse(sample_loodos_ba_elements)


# ─────────────────────────────────────────────
# TC XLSX Fixtures — Adım 6
# ─────────────────────────────────────────────

@pytest.fixture(scope="session")
def sample_tc_xlsx() -> bytes:
    """
    Minimal ama kapsamlı bir TC XLSX dosyası oluşturur.

    İçerik:
      - Cover sheet        → meta bilgiler
      - Login sheet        → Varyant 1 (TESTCASE, PRIORTY typo, 2 veri satırı + 3 boş)
      - Sepet sheet        → Varyant 2 (TEST CASE — boşluklu)
      - pinarÜrün sheet    → Varyant 4 (EXISTANCE 2/3/4, TEST CASE ID yok)
      - Regresyon Seti     → OTOMASYON kolonu var
      - DATA sheet         → skip edilmeli
      - Revision Changes   → skip edilmeli
    """
    from openpyxl import Workbook
    from io import BytesIO

    wb = Workbook()

    # ── Cover sheet ───────────────────────────────────────────────
    ws_cover = wb.active
    ws_cover.title = "Cover"
    ws_cover["C2"] = "Document Code"; ws_cover["D2"] = "LDS-PDQA-MT-KD-1"
    ws_cover["C3"] = "Project Name";  ws_cover["D3"] = "Kahve Dünyası"
    ws_cover["C4"] = "Created By";    ws_cover["D4"] = "Salih Diken"
    ws_cover["C5"] = "Create Date";   ws_cover["D5"] = "24.04.2024"

    # ── Login sheet — Varyant 1 (TESTCASE, PRIORTY typo) ─────────
    ws_login = wb.create_sheet("Login")
    v1_headers = [
        "EXISTANCE", "DATE", "APP BUNDLE", "TEST CASE ID", "BR ID", "TR ID",
        "PRIORTY", "CHANNEL", "TESTCASE TYPE", "USER TYPE",
        "TEST AREA", "TEST SCENARIO", "TESTCASE", "TEST STEPS",
        "PRECONDITION", "TEST DATA", "EXPECTED RESULT",
    ]
    for i, h in enumerate(v1_headers, 1):
        ws_login.cell(row=1, column=i, value=h)

    # Satır 2: geçerli TC
    row1 = ["New", "20.05.2024", "Kahve Dünyası", "TC_LDS_KD_LOGIN_0001",
            None, None, "HIGH", "MOBILE", "Functional", "ALL",
            "Login Ekranı", "Telefon ile Giriş", "Login TC Adı",
            "1- Uygulama açılır\n2- Login yapılır", None, None,
            "Başarılı giriş beklenir"]
    for i, v in enumerate(row1, 1):
        ws_login.cell(row=2, column=i, value=v)

    # Satır 3: CRITICAL priority
    row2 = ["New", "20.05.2024", "Kahve Dünyası", "TC_LDS_KD_LOGIN_0002",
            None, None, "CRITICAL", "MOBILE", "UI", "ALL",
            "Login Ekranı", "Hatalı Giriş", "Hatalı Login",
            "1- Yanlış şifre girilir\n2- Hata mesajı kontrol edilir", None, None,
            "Hata mesajı gösterilir"]
    for i, v in enumerate(row2, 1):
        ws_login.cell(row=3, column=i, value=v)

    # Satır 4–6: tamamen boş (filtre edilmeli)
    # (hiçbir şey yazılmaz → boş kalır)

    # ── Sepet sheet — Varyant 2 (TEST CASE boşluklu) ─────────────
    ws_sepet = wb.create_sheet("Sepet")
    v2_headers = [
        "EXISTANCE", "DATE", "APP BUNDLE", "TEST CASE ID", "BR ID", "TR ID",
        "PRIORTY", "CHANNEL", "TESTCASE TYPE", "USER TYPE",
        "TEST AREA", "TEST SCENARIO", "TEST CASE", "TEST STEPS",
        "PRECONDITION", "TEST DATA", "EXPECTED RESULT",
    ]
    for i, h in enumerate(v2_headers, 1):
        ws_sepet.cell(row=1, column=i, value=h)

    sepet_row = ["Existing", "21.05.2024", "Kahve Dünyası", "TC_LDS_KD_SEPET_0001",
                 None, None, "MEDIUM", "WEB", "UI", "Member",
                 "Sepet Ekranı", "Ürün Ekleme", "Sepet Testi",
                 "1- Ürün seçilir\n2- Sepete eklenir", None, None,
                 "Ürün sepette görünür"]
    for i, v in enumerate(sepet_row, 1):
        ws_sepet.cell(row=2, column=i, value=v)

    # ── pinarÜrün — Varyant 4 (EXISTANCE 2/3/4, TEST CASE ID yok) ──
    ws_pinar = wb.create_sheet("pinarÜrünDetay")
    v4_headers = [
        "EXISTANCE", "EXISTANCE 2", "EXISTANCE 3", "EXISTANCE 4",
        "DATE", "APP BUNDLE", "BR ID", "TR ID",
        "PRIORTY", "CHANNEL", "TESTCASE TYPE", "USER TYPE",
        "TEST AREA", "TEST SCENARIO", "TESTCASE", "TEST STEPS",
        "PRECONDITION", "TEST DATA", "EXPECTED RESULT",
    ]
    for i, h in enumerate(v4_headers, 1):
        ws_pinar.cell(row=1, column=i, value=h)

    pinar_row = ["New", "v1", "v2", "v3",
                 "22.05.2024", "Kahve Dünyası", None, None,
                 "LOW", "MOBILE", "Functional", "ALL",
                 "Ürün Detay", "Detay Görüntüleme", "Ürün Detay Testi",
                 "1- Ürün tıklanır\n2- Detay sayfası açılır", None, None,
                 "Detay sayfası yüklenir"]
    for i, v in enumerate(pinar_row, 1):
        ws_pinar.cell(row=2, column=i, value=v)

    # ── Regresyon Seti — OTOMASYON kolonu ────────────────────────
    ws_reg = wb.create_sheet("Regresyon Seti")
    reg_headers = [
        "EXISTANCE", "DATE", "APP BUNDLE", "TEST CASE ID", "BR ID", "TR ID",
        "PRIORTY", "CHANNEL", "TESTCASE TYPE", "USER TYPE",
        "TEST AREA", "TEST SCENARIO", "TESTCASE", "TEST STEPS",
        "PRECONDITION", "TEST DATA", "EXPECTED RESULT", "OTOMASYON",
    ]
    for i, h in enumerate(reg_headers, 1):
        ws_reg.cell(row=1, column=i, value=h)

    reg_row = ["New", "23.05.2024", "Kahve Dünyası", "TC_REG_0001",
               None, None, "HIGH", "MOBILE", "Regression", "ALL",
               "Splash", "Splash Kontrolü", "Splash Regresyon",
               "1- App açılır\n2- Splash kontrol edilir", None, None,
               "Splash görünür", "Otomasyon Hazır"]
    for i, v in enumerate(reg_row, 1):
        ws_reg.cell(row=2, column=i, value=v)

    # ── Skip sheet'ler ────────────────────────────────────────────
    wb.create_sheet("DATA")
    wb.create_sheet("Revision Changes")

    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


@pytest.fixture(scope="session")
def sample_tc_raw(sample_tc_xlsx) -> dict:
    """read_tc_xlsx() çıktısı (raw, parse edilmemiş)."""
    from pipeline.tc_xlsx_reader import read_tc_xlsx
    return read_tc_xlsx(sample_tc_xlsx)


@pytest.fixture(scope="session")
def sample_tc_parsed(sample_tc_raw) -> dict:
    """TCExcelParser().parse() çıktısı (normalize edilmiş)."""
    from pipeline.tc_xlsx_parser import TCExcelParser
    return TCExcelParser(sample_tc_raw).parse()
