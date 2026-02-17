"""
Tests for read_docx_structured() — Adım 1: Enhanced DOCX Reader
Kahve Dünyası formatı: bullet-list tabanlı, H1-H6, bold segments, hyperlinks
"""
import io
import pytest
from docx import Document
from docx.oxml import OxmlElement

from pipeline.document_reader import read_docx_structured


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make(build_fn) -> bytes:
    doc = Document()
    build_fn(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_hyperlink(para, url: str, text: str):
    part  = para.part
    r_id  = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(
        '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id', r_id
    )
    run_el = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    run_el.append(t)
    hyperlink.append(run_el)
    para._element.append(hyperlink)


def _elements(build_fn):
    return read_docx_structured(_make(build_fn))


# ---------------------------------------------------------------------------
# Heading extraction
# ---------------------------------------------------------------------------

class TestHeadingExtraction:
    def test_h1_detected(self):
        def build(doc):
            doc.add_heading('Proje Açıklaması', level=1)

        els = _elements(build)
        headings = [e for e in els if e['type'] == 'heading']
        assert len(headings) == 1
        assert headings[0]['level'] == 1
        assert headings[0]['text'] == 'Proje Açıklaması'

    def test_h2_detected(self):
        def build(doc):
            doc.add_heading('Splash', level=2)

        els = _elements(build)
        headings = [e for e in els if e['type'] == 'heading']
        assert headings[0]['level'] == 2
        assert headings[0]['text'] == 'Splash'

    def test_h3_detected(self):
        def build(doc):
            doc.add_heading('İş Akışı', level=3)

        els = _elements(build)
        headings = [e for e in els if e['type'] == 'heading']
        assert headings[0]['level'] == 3

    def test_h4_h5_detected(self):
        def build(doc):
            doc.add_heading('Level 4', level=4)
            doc.add_heading('Level 5', level=5)

        els = _elements(build)
        headings = [e for e in els if e['type'] == 'heading']
        assert len(headings) == 2
        assert headings[0]['level'] == 4
        assert headings[1]['level'] == 5

    def test_empty_heading_skipped(self):
        def build(doc):
            doc.add_heading('', level=1)
            doc.add_heading('Real Heading', level=2)

        els = _elements(build)
        headings = [e for e in els if e['type'] == 'heading']
        assert len(headings) == 1
        assert headings[0]['text'] == 'Real Heading'

    def test_multiple_headings_ordered(self):
        def build(doc):
            doc.add_heading('H1 Bölüm', level=1)
            doc.add_heading('H2 Ekran', level=2)
            doc.add_heading('H3 Alt', level=3)

        els = _elements(build)
        headings = [e for e in els if e['type'] == 'heading']
        assert [h['level'] for h in headings] == [1, 2, 3]


# ---------------------------------------------------------------------------
# List item extraction
# ---------------------------------------------------------------------------

class TestListItemExtraction:
    def test_list_items_have_level_field(self):
        """Bullet list item'lar 'level' alanı taşımalı."""
        def build(doc):
            doc.add_paragraph('Normal paragraf')

        # python-docx Document() varsayılan olarak numPr olmadan paragraf ekler.
        # Gerçek bullet list testi için numPr XML ekliyoruz.
        doc = Document()
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement as OE

        p = doc.add_paragraph('Birinci kural')
        pPr = OE('w:pPr')
        numPr = OE('w:numPr')
        ilvl = OE('w:ilvl')
        ilvl.set(qn('w:val'), '0')
        numId = OE('w:numId')
        numId.set(qn('w:val'), '1')
        numPr.append(ilvl)
        numPr.append(numId)
        pPr.append(numPr)
        p._element.insert(0, pPr)

        buf = io.BytesIO()
        doc.save(buf)
        els = read_docx_structured(buf.getvalue())

        list_items = [e for e in els if e['type'] == 'list_item']
        assert len(list_items) >= 1
        assert list_items[0]['level'] == 0

    def test_list_item_has_required_keys(self):
        """List item'lar text, bold_segments, links alanlarına sahip olmalı."""
        doc = Document()
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement as OE

        p = doc.add_paragraph('Alt detay')
        pPr = OE('w:pPr')
        numPr = OE('w:numPr')
        ilvl = OE('w:ilvl')
        ilvl.set(qn('w:val'), '1')
        numId = OE('w:numId')
        numId.set(qn('w:val'), '1')
        numPr.append(ilvl)
        numPr.append(numId)
        pPr.append(numPr)
        p._element.insert(0, pPr)

        buf = io.BytesIO()
        doc.save(buf)
        els = read_docx_structured(buf.getvalue())

        list_items = [e for e in els if e['type'] == 'list_item']
        assert len(list_items) >= 1
        li = list_items[0]
        assert 'text' in li
        assert 'bold_segments' in li
        assert 'links' in li
        assert li['level'] == 1


# ---------------------------------------------------------------------------
# Empty paragraph filtering
# ---------------------------------------------------------------------------

class TestEmptyParagraphFiltering:
    def test_empty_paragraphs_not_in_output(self):
        def build(doc):
            doc.add_paragraph('')
            doc.add_paragraph('   ')
            doc.add_paragraph('Gerçek içerik')

        els = _elements(build)
        texts = [e.get('text', '') for e in els]
        assert '' not in texts
        assert '   ' not in texts
        assert 'Gerçek içerik' in texts

    def test_mixed_empty_and_content(self):
        def build(doc):
            doc.add_heading('Başlık', level=1)
            doc.add_paragraph('')
            doc.add_paragraph('İçerik')

        els = _elements(build)
        assert len(els) == 2
        assert els[0]['type'] == 'heading'
        assert els[1]['text'] == 'İçerik'


# ---------------------------------------------------------------------------
# Hyperlink extraction
# ---------------------------------------------------------------------------

class TestHyperlinkExtraction:
    def test_hyperlink_url_extracted(self):
        url = 'https://app.lottiefiles.com/share/abc123'

        def build(doc):
            para = doc.add_paragraph('Animasyon: ')
            _add_hyperlink(para, url, 'Lottie Linki')

        els = _elements(build)
        all_links = [l for e in els for l in e.get('links', [])]
        assert url in all_links

    def test_paragraph_without_links_has_empty_list(self):
        def build(doc):
            doc.add_paragraph('Normal paragraf, link yok.')

        els = _elements(build)
        paras = [e for e in els if e['type'] == 'paragraph']
        assert paras[0]['links'] == []

    def test_multiple_hyperlinks_in_paragraph(self):
        url1 = 'https://app.lottiefiles.com/share/aaa'
        url2 = 'https://docs.google.com/document/d/xyz'

        def build(doc):
            para = doc.add_paragraph('Linkler: ')
            _add_hyperlink(para, url1, 'Lottie')
            _add_hyperlink(para, url2, 'Google Doc')

        els = _elements(build)
        all_links = [l for e in els for l in e.get('links', [])]
        assert url1 in all_links
        assert url2 in all_links

    def test_lottie_link_in_heading_paragraph(self):
        url = 'https://figma.com/file/abc/design'

        def build(doc):
            doc.add_heading('Splash', level=2)
            para = doc.add_paragraph('Tasarım: ')
            _add_hyperlink(para, url, 'Figma')

        els = _elements(build)
        paras = [e for e in els if e['type'] == 'paragraph']
        assert any(url in e.get('links', []) for e in paras)


# ---------------------------------------------------------------------------
# Bold segment extraction
# ---------------------------------------------------------------------------

class TestBoldSegmentExtraction:
    def test_bold_runs_extracted(self):
        doc = Document()
        para = doc.add_paragraph()
        run_normal = para.add_run('Normal metin ')
        run_bold   = para.add_run('BOLD_TERM')
        run_bold.bold = True

        buf = io.BytesIO()
        doc.save(buf)
        els = read_docx_structured(buf.getvalue())

        paras = [e for e in els if e['type'] == 'paragraph']
        assert len(paras) == 1
        assert 'BOLD_TERM' in paras[0]['bold_segments']

    def test_non_bold_runs_not_in_bold_segments(self):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run('Sadece normal metin')
        run.bold = False

        buf = io.BytesIO()
        doc.save(buf)
        els = read_docx_structured(buf.getvalue())

        paras = [e for e in els if e['type'] == 'paragraph']
        assert paras[0]['bold_segments'] == []

    def test_paragraph_has_bold_segments_key(self):
        def build(doc):
            doc.add_paragraph('Herhangi bir metin')

        els = _elements(build)
        paras = [e for e in els if e['type'] == 'paragraph']
        assert 'bold_segments' in paras[0]


# ---------------------------------------------------------------------------
# Table extraction
# ---------------------------------------------------------------------------

class TestTableExtraction:
    def test_table_extracted_with_headers_and_rows(self):
        def build(doc):
            t = doc.add_table(rows=3, cols=2)
            t.cell(0, 0).text = 'Sütun A'
            t.cell(0, 1).text = 'Sütun B'
            t.cell(1, 0).text = 'Veri 1'
            t.cell(1, 1).text = 'Veri 2'
            t.cell(2, 0).text = 'Veri 3'
            t.cell(2, 1).text = 'Veri 4'

        els = _elements(build)
        tables = [e for e in els if e['type'] == 'table']
        assert len(tables) == 1
        assert tables[0]['headers'] == ['Sütun A', 'Sütun B']
        assert len(tables[0]['rows']) == 2

    def test_empty_table_handled(self):
        def build(doc):
            doc.add_table(rows=1, cols=2)

        els = _elements(build)
        tables = [e for e in els if e['type'] == 'table']
        assert len(tables) == 1
        assert tables[0]['rows'] == []


# ---------------------------------------------------------------------------
# Element order preserved
# ---------------------------------------------------------------------------

class TestElementOrderPreserved:
    def test_heading_paragraph_order(self):
        def build(doc):
            doc.add_heading('Başlık', level=1)
            doc.add_paragraph('Paragraf içerik')

        els = _elements(build)
        assert els[0]['type'] == 'heading'
        assert els[1]['type'] == 'paragraph'

    def test_heading_paragraph_table_order(self):
        def build(doc):
            doc.add_heading('Başlık', level=2)
            doc.add_paragraph('Açıklama')
            t = doc.add_table(rows=2, cols=2)
            t.cell(0, 0).text = 'H1'; t.cell(0, 1).text = 'H2'
            t.cell(1, 0).text = 'R1'; t.cell(1, 1).text = 'R2'

        els = _elements(build)
        types = [e['type'] for e in els]
        assert types == ['heading', 'paragraph', 'table']

    def test_empty_document_returns_empty_list(self):
        def build(doc):
            pass

        els = _elements(build)
        assert els == []


# ---------------------------------------------------------------------------
# DocxImportOrchestrator — Adım 3 testleri
# ---------------------------------------------------------------------------

class TestDocxImportOrchestratorStats:
    """_calculate_stats() doğru sayıları döndürmeli."""

    def _stats(self, elements):
        from pipeline.docx_import_orchestrator import DocxImportOrchestrator
        return DocxImportOrchestrator()._calculate_stats(elements)

    def _h(self, level, text):
        return {"type": "heading", "level": level, "text": text}

    def _li(self, level, text, links=None):
        return {"type": "list_item", "level": level, "text": text, "links": links or []}

    def _p(self, text, links=None):
        return {"type": "paragraph", "text": text, "links": links or []}

    def _tbl(self):
        return {"type": "table", "headers": [], "rows": []}

    def test_empty_document(self):
        stats = self._stats([])
        assert stats == {"headings": 0, "screens": 0, "list_items": 0,
                         "paragraphs": 0, "tables": 0, "links": 0}

    def test_heading_counts(self):
        elements = [self._h(1, "H1"), self._h(2, "H2a"), self._h(2, "H2b"), self._h(3, "H3")]
        stats = self._stats(elements)
        assert stats["headings"] == 4
        assert stats["screens"] == 2  # sadece H2 sayılır

    def test_list_item_count(self):
        elements = [self._li(0, "A"), self._li(1, "B"), self._li(2, "C")]
        stats = self._stats(elements)
        assert stats["list_items"] == 3

    def test_table_count(self):
        elements = [self._tbl(), self._tbl()]
        stats = self._stats(elements)
        assert stats["tables"] == 2

    def test_link_count(self):
        url1 = "https://figma.com/a"
        url2 = "https://lottiefiles.com/b"
        elements = [
            self._p("Para", links=[url1]),
            self._li(0, "Item", links=[url2]),
        ]
        stats = self._stats(elements)
        assert stats["links"] == 2


class TestDocxImportOrchestratorDetect:
    """_detect_template() şablon tipini doğru algılamalı."""

    def _detect(self, elements, stats=None):
        from pipeline.docx_import_orchestrator import DocxImportOrchestrator
        orch = DocxImportOrchestrator()
        if stats is None:
            stats = orch._calculate_stats(elements)
        return orch._detect_template(elements, stats)

    def _h(self, level, text):
        return {"type": "heading", "level": level, "text": text}

    def _li(self, level, text):
        return {"type": "list_item", "level": level, "text": text, "links": []}

    def _tbl(self):
        return {"type": "table", "headers": [], "rows": []}

    def test_loodos_ba_bullet_detected_with_2_markers(self):
        elements = [
            self._h(3, "Açıklama"),
            self._h(3, "İş Akışı"),
            self._li(0, "Kural 1"),
            self._li(0, "Kural 2"),
        ]
        assert self._detect(elements) == "loodos_ba_bullet"

    def test_loodos_ba_bullet_detected_with_all_3_markers(self):
        elements = [
            self._h(3, "Açıklama"),
            self._h(3, "İş Akışı"),
            self._h(3, "Tasarım Dosyaları"),
            self._li(0, "Kural"),
        ]
        assert self._detect(elements) == "loodos_ba_bullet"

    def test_loodos_ba_bullet_case_insensitive(self):
        """Başlıklar küçük harfle de eşleşmeli."""
        elements = [
            self._h(3, "açıklama"),
            self._h(3, "iş akışı"),
            self._li(0, "Kural"),
        ]
        assert self._detect(elements) == "loodos_ba_bullet"

    def test_one_marker_not_enough_for_bullet(self):
        """Tek marker → loodos_ba_bullet değil."""
        elements = [
            self._h(3, "İş Akışı"),
            self._li(0, "Kural"),
        ]
        result = self._detect(elements)
        assert result != "loodos_ba_bullet"

    def test_loodos_ba_table_detected_with_many_tables(self):
        elements = [self._tbl() for _ in range(5)]
        stats = {"headings": 0, "screens": 0, "list_items": 0,
                 "paragraphs": 0, "tables": 5, "links": 0}
        assert self._detect(elements, stats) == "loodos_ba_table"

    def test_3_tables_not_enough_for_table_template(self):
        stats = {"tables": 3, "list_items": 0}
        assert self._detect([], stats) != "loodos_ba_table"

    def test_generic_when_no_markers(self):
        elements = [
            self._h(1, "Giriş"),
            self._h(2, "Bölüm 1"),
        ]
        assert self._detect(elements) == "generic"

    def test_bullet_preferred_over_table_when_list_items_dominate(self):
        """2 Loodos marker + list_items > tables → bullet seçilmeli."""
        elements = [
            self._h(3, "Açıklama"),
            self._h(3, "İş Akışı"),
            self._li(0, "K1"),
            self._li(0, "K2"),
            self._tbl(),
        ]
        assert self._detect(elements) == "loodos_ba_bullet"


class TestDocxImportOrchestratorConfidence:
    """_calculate_confidence() doğru skor hesaplamalı."""

    def _conf(self, ekranlar, stats=None):
        from pipeline.docx_import_orchestrator import DocxImportOrchestrator
        content_json = {"ekranlar": ekranlar}
        s = stats or {"links": 0, "tables": 0}
        return DocxImportOrchestrator()._calculate_confidence(content_json, s)

    def _screen(self, name, has_rules=True):
        kurallar = [{"kural": "Kural", "level": 0, "alt_detaylar": [],
                     "bold_refs": [], "links": []}] if has_rules else []
        return {
            "ekran_adi": name,
            "aciklama": "",
            "tasarim_dosyalari": [],
            "is_akislari": [{"baslik": "İş Akışı", "kurallar": kurallar}],
        }

    def test_no_screens_zero_confidence(self):
        assert self._conf([]) == 0.0

    def test_one_screen_no_rules(self):
        screen = {
            "ekran_adi": "X",
            "is_akislari": [{"baslik": "İş Akışı", "kurallar": []}],
        }
        score = self._conf([screen])
        assert score == pytest.approx(0.3)  # sadece "en az 1 ekran" puanı

    def test_one_screen_with_rules(self):
        score = self._conf([self._screen("Splash")])
        # 0.3 (ekran var) + 0.3 (kural var) + 0.2 (oran=1.0)
        assert score == pytest.approx(0.8)

    def test_five_screens_all_with_rules(self):
        screens = [self._screen(f"Ekran{i}") for i in range(5)]
        score = self._conf(screens)
        # 0.3 + 0.2 (>=5) + 0.3 + 0.2 = 1.0
        assert score == pytest.approx(1.0)

    def test_partial_rules_reduces_score(self):
        """5 ekranın 2'sinde kural var → oran=0.4."""
        screens = [self._screen(f"E{i}", has_rules=(i < 2)) for i in range(5)]
        score = self._conf(screens)
        # 0.3 + 0.2 + 0.3 + 0.4*0.2 = 0.88
        assert score == pytest.approx(0.88)

    def test_confidence_capped_at_1(self):
        screens = [self._screen(f"E{i}") for i in range(10)]
        assert self._conf(screens) <= 1.0


class TestDocxImportOrchestratorWarnings:
    """_generate_warnings() doğru uyarıları üretmeli."""

    def _warn(self, ekranlar, links=0):
        from pipeline.docx_import_orchestrator import DocxImportOrchestrator
        content_json = {"ekranlar": ekranlar}
        stats = {"links": links}
        return DocxImportOrchestrator()._generate_warnings(content_json, stats)

    def _screen(self, name, has_rules=True):
        kurallar = [{"kural": "K"}] if has_rules else []
        return {
            "ekran_adi": name,
            "is_akislari": [{"baslik": "İş Akışı", "kurallar": kurallar}],
        }

    def test_no_warnings_when_all_ok(self):
        warnings = self._warn([self._screen("Splash")], links=3)
        assert warnings == []

    def test_warning_for_screen_without_rules(self):
        warnings = self._warn([self._screen("OTP", has_rules=False)])
        assert any("OTP" in w for w in warnings)

    def test_warning_for_no_links(self):
        warnings = self._warn([self._screen("Splash")], links=0)
        assert any("link" in w.lower() for w in warnings)

    def test_multiple_screens_multiple_warnings(self):
        screens = [
            self._screen("A", has_rules=False),
            self._screen("B", has_rules=True),
            self._screen("C", has_rules=False),
        ]
        warnings = self._warn(screens, links=0)
        assert len(warnings) == 3  # 2 ekran + 1 link

    def test_screen_with_empty_is_akislari(self):
        """is_akislari boş olan ekran da uyarı almalı."""
        screen = {"ekran_adi": "Boş Ekran", "is_akislari": []}
        warnings = self._warn([screen])
        assert any("Boş Ekran" in w for w in warnings)


class TestDocxImportOrchestratorIntegration:
    """import_docx() uçtan uca DOCX bytes ile çalışmalı."""

    def _make_loodos_docx(self) -> bytes:
        """Loodos BA bullet formatında test DOCX oluştur."""
        doc = Document()
        doc.add_heading('Mobil Uygulama Gereksinimleri', level=1)
        doc.add_heading('Splash', level=2)
        doc.add_heading('Açıklama', level=3)
        doc.add_paragraph('Splash ekranı.')
        doc.add_heading('İş Akışı', level=3)
        # Bullet item için numPr ekliyoruz
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        for text in ['Uygulama açılır.', 'Logo gösterilir.']:
            p = doc.add_paragraph(text)
            pPr = OxmlElement('w:pPr')
            numPr = OxmlElement('w:numPr')
            ilvl = OxmlElement('w:ilvl')
            ilvl.set(qn('w:val'), '0')
            numId = OxmlElement('w:numId')
            numId.set(qn('w:val'), '1')
            numPr.append(ilvl); numPr.append(numId)
            pPr.append(numPr)
            p._element.insert(0, pPr)
        doc.add_heading('Login', level=2)
        doc.add_heading('Tasarım Dosyaları', level=3)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def test_import_returns_required_keys(self):
        from pipeline.docx_import_orchestrator import DocxImportOrchestrator
        result = DocxImportOrchestrator().import_docx(self._make_loodos_docx())
        for key in ('success', 'doc_type', 'template', 'content_json',
                    'confidence', 'stats', 'warnings'):
            assert key in result

    def test_loodos_ba_bullet_detected_end_to_end(self):
        from pipeline.docx_import_orchestrator import DocxImportOrchestrator
        result = DocxImportOrchestrator().import_docx(self._make_loodos_docx())
        assert result['template'] == 'loodos_ba_bullet'

    def test_screens_in_content_json(self):
        from pipeline.docx_import_orchestrator import DocxImportOrchestrator
        result = DocxImportOrchestrator().import_docx(self._make_loodos_docx())
        screens = result['content_json'].get('ekranlar', [])
        names = [s['ekran_adi'] for s in screens]
        assert 'Splash' in names
        assert 'Login' in names

    def test_success_true_when_confidence_above_threshold(self):
        from pipeline.docx_import_orchestrator import DocxImportOrchestrator
        result = DocxImportOrchestrator().import_docx(self._make_loodos_docx())
        assert result['success'] is True

    def test_stats_populated(self):
        from pipeline.docx_import_orchestrator import DocxImportOrchestrator
        result = DocxImportOrchestrator().import_docx(self._make_loodos_docx())
        stats = result['stats']
        assert stats['screens'] >= 2
        assert stats['headings'] > 0


# ---------------------------------------------------------------------------
# Fixture tabanlı testler — conftest.py sample_loodos_ba_docx kullanır
# ---------------------------------------------------------------------------

class TestWithFixtures:
    """
    conftest.py'deki sample_loodos_ba_docx fixture'ını kullanan testler.
    Gerçek doküman yapısını doğrular: H1 meta, H2 ekranlar, H3 alt bölümler,
    nested bullet list, hyperlink'ler, bold run'lar.
    """

    def test_fixture_produces_bytes(self, sample_loodos_ba_docx):
        """Fixture geçerli bytes döndürmeli."""
        assert isinstance(sample_loodos_ba_docx, bytes)
        assert len(sample_loodos_ba_docx) > 0
        # DOCX (ZIP) formatı PK ile başlar
        assert sample_loodos_ba_docx[:2] == b'PK'

    def test_fixture_elements_has_headings(self, sample_loodos_ba_elements):
        """Element listesinde heading'ler olmalı."""
        headings = [e for e in sample_loodos_ba_elements if e['type'] == 'heading']
        assert len(headings) > 0

    def test_fixture_elements_has_list_items(self, sample_loodos_ba_elements):
        """Element listesinde list item'lar olmalı."""
        list_items = [e for e in sample_loodos_ba_elements if e['type'] == 'list_item']
        assert len(list_items) > 0

    def test_fixture_h2_screens_present(self, sample_loodos_ba_elements):
        """Splash ve Login H2 heading'leri mevcut olmalı."""
        h2_texts = [e['text'] for e in sample_loodos_ba_elements
                    if e['type'] == 'heading' and e['level'] == 2]
        assert 'Splash' in h2_texts
        assert 'Login' in h2_texts

    def test_fixture_h3_subsections_present(self, sample_loodos_ba_elements):
        """İş Akışı ve Açıklama H3 başlıkları mevcut olmalı."""
        h3_texts = {e['text'] for e in sample_loodos_ba_elements
                    if e['type'] == 'heading' and e['level'] == 3}
        assert 'İş Akışı' in h3_texts
        assert 'Açıklama' in h3_texts
        assert 'Tasarım Dosyaları' in h3_texts

    def test_fixture_hyperlinks_present(self, sample_loodos_ba_elements):
        """Figma ve Lottie linkleri element listesinde bulunmalı."""
        all_links = [url for e in sample_loodos_ba_elements for url in e.get('links', [])]
        assert any('figma.com' in url for url in all_links)
        assert any('lottiefiles.com' in url for url in all_links)

    def test_fixture_parsed_screens(self, sample_loodos_ba_parsed):
        """Parse sonucunda Splash ve Login ekranları çıkarılmalı."""
        names = [e['ekran_adi'] for e in sample_loodos_ba_parsed['ekranlar']]
        assert 'Splash' in names
        assert 'Login' in names

    def test_fixture_parsed_is_akislari(self, sample_loodos_ba_parsed):
        """Her ekranın is_akislari listesi dolu olmalı."""
        for ekran in sample_loodos_ba_parsed['ekranlar']:
            assert len(ekran['is_akislari']) >= 1

    def test_fixture_parsed_rule_tree_nested(self, sample_loodos_ba_parsed):
        """Splash iş akışı L0 → L1 nesting içermeli."""
        splash = next(e for e in sample_loodos_ba_parsed['ekranlar']
                      if e['ekran_adi'] == 'Splash')
        ia = splash['is_akislari'][0]
        l0_with_children = [k for k in ia['kurallar'] if k['alt_detaylar']]
        assert len(l0_with_children) >= 1

    def test_fixture_parsed_meta(self, sample_loodos_ba_parsed):
        """Meta bilgileri çıkarılmalı: platform ve proje açıklaması."""
        meta = sample_loodos_ba_parsed['meta']
        assert meta['proje_aciklamasi'] != ''
        assert meta['platform'] != ''

    def test_fixture_parsed_links_categorized(self, sample_loodos_ba_parsed):
        """Figma ve Lottie linkleri doğru kategorilere gitmiş olmalı."""
        linkler = sample_loodos_ba_parsed['linkler']
        assert len(linkler['figma']) >= 1
        assert len(linkler['lottie']) >= 1

    def test_fixture_orchestrator_end_to_end(self, sample_loodos_ba_docx):
        """Orchestrator fixture DOCX üzerinde başarıyla çalışmalı."""
        from pipeline.docx_import_orchestrator import DocxImportOrchestrator
        result = DocxImportOrchestrator().import_docx(sample_loodos_ba_docx)
        assert result['success'] is True
        assert result['template'] == 'loodos_ba_bullet'
        assert result['confidence'] > 0.5
        assert result['stats']['screens'] >= 2
